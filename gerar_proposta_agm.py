#!/usr/bin/env python3
"""Gera proposta estratégica AGM Distribuidora (DOCX para Google Docs)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND = RGBColor(0xC4, 0x10, 0x20)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
LINE = "D0D0D0"


def set_run(run, *, size=11, bold=False, color=INK, font="Arial"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rPr.insert(0, rFonts)


def add_para(doc, text, *, size=11, bold=False, color=INK, space_after=8, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_rich(doc, parts, *, space_after=8, space_before=0):
    """parts = list of (text, bold, color, size)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, bold, color, size in parts:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    add_para(doc, text, size=22, bold=True, color=BRAND, space_after=6, space_before=0)


def h2(doc, text):
    add_para(doc, text, size=14, bold=True, color=BRAND, space_after=6, space_before=16)


def h3(doc, text):
    add_para(doc, text, size=12, bold=True, color=INK, space_after=4, space_before=10)


def body(doc, text, **kw):
    add_para(doc, text, size=11, color=INK, space_after=8, **kw)


def muted(doc, text):
    add_para(doc, text, size=10, color=MUTED, space_after=10)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text)
        set_run(r2, size=11, bold=False, color=INK)
    else:
        r = p.add_run(text)
        set_run(r, size=11, bold=False, color=INK)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C41020")
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10, color=INK, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), center=True)
        shade_cell(cell, "C41020")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, bold=(c_idx == 0), size=10)
            if r_idx % 2 == 1:
                shade_cell(cell, "F7F7F7")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # CAPA
    add_para(doc, "V4 COMPANY", size=11, bold=True, color=BRAND, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "PROPOSTA DE PARCERIA",
        size=20,
        bold=True,
        color=INK,
        space_after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "AGM Distribuidora",
        size=16,
        bold=True,
        color=BRAND,
        space_after=6,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    muted(doc, "Assessoria de Growth · Execução do plano da Estruturação Estratégica · 12 meses")
    add_para(
        doc,
        "Documento estratégico de continuidade | Julho/2026",
        size=10,
        color=MUTED,
        space_after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_hr(doc)

    # 1. CONTEXTO
    h2(doc, "1. Contexto da parceria")
    body(
        doc,
        "A AGM Distribuidora construiu uma operação sólida no offline: base B2B recorrente, "
        "vendedores experientes, logística por rotas e um B2B (Mercos) que já funciona como "
        "canal relevante de pedidos. O desafio atual não é falta de mercado — é transformar o "
        "digital em um canal previsível de aquisição e recompra, capaz de sustentar a meta de "
        "crescimento (de aproximadamente R$ 330 mil/mês para R$ 500 mil/mês).",
    )
    body(
        doc,
        "Ao longo da Estruturação Estratégica (E.E.), mapeamos mercado, concorrência, personas, "
        "funil, oferta, tracking, redes, tráfego e oportunidades de reativação da base. "
        "O diagnóstico confirma: existe demanda digital já validada (clientes chegando por "
        "Instagram inclusive de outras cidades), porém sem mensuração, sem processo de "
        "atendimento inbound e sem rotina comercial digital integrada.",
    )
    body(
        doc,
        "Este documento apresenta a proposta de continuidade: colocar o plano da E.E. em "
        "operação, com time dedicado da V4, pelo período de 12 meses.",
    )

    # 2. O QUE JÁ FOI CONSTRUÍDO
    h2(doc, "2. O que a Estruturação Estratégica já entregou")
    body(doc, "A E.E. organizou a base para a tomada de decisão e para a execução. Entre os principais entregáveis e direcionamentos:")
    bullet(doc, "Diagnóstico de marketing, vendas e canais digitais da AGM.")
    bullet(doc, "Leitura de mercado, concorrentes e oportunidades regionais (incluindo potencial local sem expansão territorial).")
    bullet(doc, "Definição de personas prioritárias: lojista de bairro, lojista sazonal e base inativa.")
    bullet(doc, "Matriz de ofertas e funil (cadastro → primeiro pedido → cliente ativo → recompra).")
    bullet(doc, "Direcionamento de comunicação (bio, destaques, manuais de copy e identidade visual).")
    bullet(doc, "Plano de tráfego (Meta + Google) e projeção de 12 meses.")
    bullet(doc, "Prioridades de retenção/reativação (semáforo de base e réguas de contato).")
    bullet(doc, "Roadmap de fundação → ativação → consolidação (30 a 90 dias e além).")
    body(
        doc,
        "Com o mapa definido, o próximo passo natural é a operação assistida e contínua — "
        "para que estratégia vire rotina, lead e receita.",
    )

    # 3. PROPOSTA
    h2(doc, "3. Proposta: Assessoria de Growth (Executar)")
    body(
        doc,
        "A Assessoria de Growth é o braço de execução da V4. Em vez de deixar o plano dependente "
        "da agenda interna já sobrecarregada da AGM, a V4 assume a operação digital com um time "
        "multidisciplinar, conectado ao comercial e aos objetivos de faturamento.",
    )
    h3(doc, "Objetivo da parceria")
    bullet(doc, "Tornar o digital um canal estruturado de geração de demanda B2B.")
    bullet(doc, "Aumentar a base de clientes ativos e acelerar o primeiro pedido.")
    bullet(doc, "Organizar recompra e reativação com processo (não só esforço espontâneo).")
    bullet(doc, "Gerar previsibilidade: origem do lead, custo, conversão e retorno.")
    bullet(doc, "Reduzir a dependência exclusiva do esforço manual do sócio no marketing.")

    h3(doc, "Escopo de execução (12 meses)")
    bullet(doc, "Gestão e otimização de tráfego pago (Meta Ads e Google Ads), com ritmo de investimento alinhado à AGM.")
    bullet(doc, "Produção de copy e criativos orientados a performance (não apenas engajamento).")
    bullet(doc, "Social media com calendário e comunicação alinhada ao funil B2B.")
    bullet(doc, "Landing pages / pontos de captura para rastrear origem e qualificar entrada de leads.")
    bullet(doc, "Acompanhamento de performance com rotinas claras de alinhamento.")
    bullet(doc, "Integração do plano de aquisição com atendimento e recompra.")

    # 4. TIME
    h2(doc, "4. Time dedicado e como funciona o dia a dia")
    body(doc, "A AGM passa a contar com um time V4 responsável por operar o plano:")
    add_table(
        doc,
        ["Papel", "Responsabilidade"],
        [
            ["Account Manager", "Orquestra a conta, prioriza o backlog, alinha expectativas e garante ritmo de entrega."],
            ["Copywriter", "Mensagens, anúncios, roteiros e comunicação orientada a conversão B2B."],
            ["Designer", "Criativos, peças e padronização visual para campanhas e redes."],
            ["Gestor de Tráfego (GT)", "Planeja, veicula e otimiza campanhas em Meta e Google."],
            ["Social Media", "Presença contínua no Instagram/redes com foco em demanda e relacionamento."],
        ],
        col_widths=[1.8, 4.8],
    )
    h3(doc, "Rotinas de alinhamento")
    bullet(doc, "Contato contínuo via grupo/canal da conta (acompanhamento do dia a dia).")
    bullet(doc, "Reuniões quinzenais de performance (números, aprendizados e próximos testes).")
    bullet(doc, "Reuniões estratégicas trimestrais (visão de ciclo, metas e ajustes de rota).")
    body(
        doc,
        "Papel da AGM na operação: aprovar direcionamentos, liberar acessos, validar criativos "
        "quando necessário e manter o time comercial alinhado ao fluxo de leads. "
        "A operação digital fica sob responsabilidade da V4.",
    )

    # 5. KOMMO
    h2(doc, "5. Implementação do CRM Kommo (bonificada)")
    body(
        doc,
        "Para que aquisição não “vaze” no atendimento, a proposta inclui a implementação do "
        "Kommo CRM de forma bonificada na parceria de Growth.",
    )
    body(doc, "Por que isso importa para a AGM:")
    bullet(doc, "Centralizar leads vindos de Instagram, campanhas, indicação e outros canais.")
    bullet(doc, "Organizar o primeiro contato com SLA (atendimento rápido e rastreável).")
    bullet(doc, "Criar visão de funil: novo lead → atendimento → primeiro pedido → recompra.")
    bullet(doc, "Sustentar réguas de reativação da base (clientes inativos e cadastrados sem compra).")
    bullet(doc, "Dar visibilidade ao Marcelo e à Débora sobre o que está entrando e o que está avançando.")
    body(
        doc,
        "Observação: a ferramenta só gera resultado com uso disciplinado. Por isso a "
        "implementação vem acompanhada de configuração orientada ao processo da AGM e "
        "alinhamento com o time que vai operar o dia a dia comercial.",
    )

    # 6. ESTRUTURAÇÃO COMERCIAL (complementar)
    h2(doc, "6. Frente complementar: Estruturação Comercial")
    body(
        doc,
        "Além da Assessoria de Growth, a V4 recomenda uma Estruturação Comercial — um sprint "
        "para organizar processos, scripts, playbooks e treinamento do time, preparando a "
        "máquina comercial para receber e converter a demanda digital com mais consistência.",
    )
    body(doc, "Essa frente aborda, entre outros pontos:")
    bullet(doc, "Mapeamento e desenho do processo comercial (do lead ao pedido).")
    bullet(doc, "Roteiros e playbooks de atendimento (WhatsApp e time interno/representantes).")
    bullet(doc, "Treinamento prático do time para leads inbound.")
    bullet(doc, "Alinhamento de metas e rituais comerciais ligados ao funil digital.")
    body(
        doc,
        "Growth gera e qualifica demanda. Estruturação Comercial fortalece a conversão e a "
        "recompra. Juntas, as frentes reduzem o risco de “trazer lead e não converter”.",
    )

    # 7. PRIORIDADES 90 DIAS
    h2(doc, "7. Prioridades dos primeiros 90 dias")
    add_table(
        doc,
        ["Fase", "Foco", "Resultado esperado"],
        [
            [
                "Fundação",
                "Tracking, acessos, Kommo, pontos de captura, alinhamento de atendimento",
                "Base mensurável e processo mínimo de entrada de leads",
            ],
            [
                "Ativação",
                "Tráfego pago, calendário social, criativos, ritmo de geração",
                "Motor digital ligado com aprendizado semanal",
            ],
            [
                "Consolidação",
                "Otimização de campanhas + reativação/recompra da base",
                "Previsibilidade crescente e melhor aproveitamento da carteira",
            ],
        ],
        col_widths=[1.4, 2.6, 2.6],
    )
    body(
        doc,
        "As prioridades respeitam a realidade operacional da AGM (correções de catálogo/site, "
        "sazonalidade e time enxuto): começamos pelo que desbloqueia resultado com menos "
        "fricção e evoluímos em ciclos.",
    )

    # 8. INVESTIMENTO
    h2(doc, "8. Investimento")
    body(
        doc,
        "Valores da proposta apresentada na reunião de continuidade da Estruturação Estratégica:",
    )

    h3(doc, "A) Assessoria de Growth — Executar (12 meses)")
    add_table(
        doc,
        ["Item", "Condição", "Valor"],
        [
            ["Mensalidade", "12 meses de operação", "R$ 6.172,87 / mês"],
            ["Período", "Contrato anual operacional", "12 meses"],
            ["Inclui", "Time (Account, Copy, Design, GT e Social) + operação do plano", "Incluso"],
            ["CRM Kommo", "Implementação na parceria", "Bonificado"],
        ],
        col_widths=[2.0, 3.0, 1.6],
    )
    add_rich(
        doc,
        [
            ("Referência de investimento anual da Assessoria: ", False, INK, 11),
            ("R$ 74.074,44", True, BRAND, 11),
            (" (12 × R$ 6.172,87).", False, INK, 11),
        ],
        space_after=10,
    )

    h3(doc, "B) Estruturação Comercial (sprint complementar)")
    add_table(
        doc,
        ["Item", "Condição", "Valor"],
        [
            ["Projeto", "Estruturação comercial / processos e treinamento", "R$ 11.523,21"],
            ["Parcelamento", "12 vezes", "12 × R$ 960,00"],
        ],
        col_widths=[2.0, 3.0, 1.6],
    )

    h3(doc, "Composição sugerida")
    body(
        doc,
        "Recomendação estratégica: iniciar a Assessoria de Growth para colocar o plano em "
        "execução imediatamente, com Kommo bonificado, e avançar a Estruturação Comercial "
        "para fortalecer conversão e rotina do time comercial em paralelo ou em sequência "
        "curta — conforme capacidade de caixa e prioridade da diretoria.",
    )

    # 9. PRÓXIMOS PASSOS
    h2(doc, "9. Próximos passos")
    bullet(doc, "Validação da proposta com Marcelo e Débora (diretoria / financeiro).")
    bullet(doc, "Assinatura do contrato da frente escolhida (Growth e/ou Estruturação Comercial).")
    bullet(doc, "Setup interno V4 em até 7 dias úteis após assinatura (acessos, repasse estratégico e preparação do time).")
    bullet(doc, "Kickoff oficial com apresentação do time e início do primeiro ciclo de execução.")
    body(
        doc,
        "A V4 permanece à disposição para esclarecer qualquer ponto desta proposta na "
        "devolutiva, com foco em tornar a decisão simples, segura e conectada ao objetivo "
        "de crescimento da AGM.",
    )

    add_hr(doc)
    add_para(
        doc,
        "V4 Company  ·  Assessoria em Marketing e Vendas",
        size=10,
        bold=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_para(
        doc,
        "AGM Distribuidora — Proposta de continuidade pós Estruturação Estratégica",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )

    out = "/workspace/Proposta-AGM-Distribuidora-Assessoria-Growth.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
