#!/usr/bin/env python3
"""Gera a proposta Work Bank CRM Vendas para abrir no Word (.html Word + .docx)."""

from html import escape
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RED = "A10F14"
RED2 = "E50914"
INK = "1A1A1A"
MUT = "4A4A4A"
CREAM = "F7F5F4"
LINE = "E6D4D4"
WHITE = "FFFFFF"
SOFT = "F8EFEF"
DARK = "280001"

OUT_HTML = Path("/workspace/proposta-workbank-crm-vendas-word.html")
OUT_DOCX = Path("/workspace/proposta-workbank-crm-vendas.docx")
OUT_DOC = Path("/workspace/proposta-workbank-crm-vendas.doc")


def set_run_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="E6D4D4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_width(table, width_cm=17):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def cell_text(cell, text, *, bold=False, size=11, color=INK, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    set_cell_margins(cell)
    set_cell_borders(cell)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 12.5, bold=True, color=RED if level == 1 else RED2)
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=italic, color=MUT if italic else INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11, color=MUT)


def grid(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_text(cell, h, bold=True, size=9, color=WHITE)
        shade_cell(cell, RED)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell_text(cell, val, bold=(c == 0), size=10, color=INK if c == 0 else MUT)
            if r % 2 == 0:
                shade_cell(cell, SOFT)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def cover_table(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.text = ""
    shade_cell(cell, DARK)
    set_cell_margins(cell, 200, 200, 200, 200)
    p = cell.paragraphs[0]
    r = p.add_run("PROPOSTA COMERCIAL")
    set_run_font(r, size=9, bold=True, color="F5B5B8")
    p2 = cell.add_paragraph()
    r2 = p2.add_run("CRM Vendas")
    set_run_font(r2, size=26, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    r3 = p3.add_run("Work Bank  ×  V4 Company")
    set_run_font(r3, size=14, color="FFD0D0")
    p4 = cell.add_paragraph()
    r4 = p4.add_run(
        "Fundação de CRM Vendas em 3 semanas — pipeline de crédito, templates e rotina "
        "para o time que entra em setembro."
    )
    set_run_font(r4, size=11, color="E8D0D0")
    doc.add_paragraph()


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.text = ""
    shade_cell(cell, SOFT)
    set_cell_margins(cell, 140, 140, 160, 160)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    set_run_font(r, size=8.5, bold=True, color=RED2)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, color=INK)
    doc.add_paragraph()


def price_block(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.text = ""
    right.text = ""
    shade_cell(left, DARK)
    shade_cell(right, SOFT)
    set_cell_margins(left, 180, 180, 180, 180)
    set_cell_margins(right, 180, 180, 180, 180)
    p = left.paragraphs[0]
    r = p.add_run("TABELA · IMPLEMENTAÇÃO CRM VENDAS")
    set_run_font(r, size=9, bold=True, color="F5B5B8")
    p2 = left.add_paragraph()
    r2 = p2.add_run("R$ 9.145,16")
    set_run_font(r2, size=14, color="C9A0A0")
    r2.font.strike = True
    p3 = left.add_paragraph()
    r3 = p3.add_run("CONDIÇÃO WORK BANK")
    set_run_font(r3, size=9, bold=True, color="F5B5B8")
    p4 = left.add_paragraph()
    r4 = p4.add_run("R$ 5.639,21")
    set_run_font(r4, size=26, bold=True, color=WHITE)
    p5 = left.add_paragraph()
    r5 = p5.add_run("3 semanas · CRM Vendas Basic · pipeline, templates, treino e documentação")
    set_run_font(r5, size=11, color="E8D0D0")

    q = right.paragraphs[0]
    rq = q.add_run("O QUE ESTÁ INCLUSO")
    set_run_font(rq, size=8.5, bold=True, color=RED2)
    for title, body in (
        ("Funil instalado", "Pipeline de crédito, campos, base, 3 templates e o time treinado no uso diário."),
        ("Por que agora", "O closer novo herda processo na ferramenta, não no WhatsApp."),
        ("Próximo passo", "Briefing + acessos. Em 3 semanas o funil está no ar."),
    ):
        p = right.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, size=12, bold=True, color=INK)
        p2 = right.add_paragraph()
        run2 = p2.add_run(body)
        set_run_font(run2, size=10, color=MUT)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2)

    cover_table(doc)
    grid(
        doc,
        ["Campo", "Informação"],
        [
            ["Cliente", "Work Bank (Marchi Soluções Financeiras)"],
            ["Produto", "Implementação CRM Vendas Basic"],
            ["Prazo", "3 semanas  ·  ~13,5 horas"],
            ["Investimento", "R$ 5.639,21  (tabela: R$ 9.145,16)"],
            ["Condução", "V4 Company  ·  Account Manager + Profissional de CRM + Copy (templates)"],
            ["Momento", "Time comercial novo entra em setembro"],
        ],
        [4.5, 12.5],
    )
    callout(
        doc,
        "LEITURA CENTRAL",
        "A Work Bank não tem problema de demanda. Tem um comercial artesanal: 90% indicação, "
        "funil na cabeça do closer e time novo sem processo na ferramenta. Esta proposta instala "
        "o esqueleto do CRM de Vendas — um pipeline de crédito, templates e rotina — em 3 semanas.",
    )

    add_heading(doc, "1. Objetivo")
    add_body(
        doc,
        "Instalar a fundação comercial da Work Bank no CRM de Vendas: um pipeline mapeado no ciclo "
        "de crédito, campos essenciais, usuários do time, importação da base, 3 templates de e-mail "
        "SDR/Closer, treinamento básico e documentação da configuração. O resultado esperado não é "
        "um manual na gaveta. É o closer cadastrando oportunidade e o gestor enxergando o funil no dia 1.",
    )

    add_heading(doc, "2. O momento Work Bank")
    add_body(
        doc,
        "Assessoria de crédito (B2B e B2C), mais de 80 instituições, atuação em Minas Gerais e entorno. "
        "Produtos: home equity, garantia de veículo, capital de giro e consórcios. Receita atual de "
        "cerca de R$ 80 mil/mês, com ambição de R$ 20 milhões em desembolso.",
    )
    grid(
        doc,
        ["Dado", "Leitura"],
        [
            ["R$ 80 mil/mês", "Modelo validado. A máquina ainda é artesanal."],
            ["90% indicação", "Networking segura a operação. Sem funil, a previsão some."],
            ["Setembro", "Time comercial novo entra. Sem processo no CRM, cada um inventa o próprio."],
            ["R$ 20 milhões", "Ambição de desembolso. O closer precisa de pipeline no dia 1."],
        ],
        [5, 12],
    )

    add_heading(doc, "3. Diagnóstico")
    add_body(
        doc,
        "Três vazamentos travam a previsibilidade. Começar errado — CRM como bloco de notas — faz o "
        "time de setembro herdar o vício.",
    )
    grid(
        doc,
        ["Vazamento", "Efeito"],
        [
            ["Oportunidade sem dono", "Indicação entra no WhatsApp. Ninguém registra. Follow-up depende de memória."],
            ["Funil na cabeça", "Cada closer vende de um jeito. Gestor não confia no número. Win rate vira opinião."],
            ["Time novo sem rampa", "Sem processo no sistema, onboarding pode ir a 4 meses. O conhecimento sai com a pessoa."],
        ],
        [5.2, 11.8],
    )
    callout(
        doc,
        "O RISCO",
        "A indicação existe e o crédito não fecha. Sem o 1º pipeline, o gestor não sabe por que o "
        "crédito morreu: simulação, parceiro ou documentação.",
    )

    add_heading(doc, "4. O que é o CRM de Vendas")
    add_body(
        doc,
        "Não é agenda de contato. É o esqueleto do processo comercial: o lugar onde cada pedido de "
        "crédito vive — do primeiro “me indica” até o desembolso. A metodologia sai da cabeça do "
        "vendedor e vira rotina na ferramenta.",
    )
    grid(
        doc,
        ["Peça", "Para que serve"],
        [
            ["Dono", "Toda oportunidade tem responsável."],
            ["Etapa", "O crédito anda no funil, não no chat."],
            ["Histórico", "O gestor vê o funil, não pergunta um a um."],
            ["Rotina", "O closer novo herda o processo no dia 1."],
        ],
        [4, 13],
    )

    add_heading(doc, "5. Case — XP Inc. × V4 Exclusive Experts")
    add_body(
        doc,
        "Em serviços financeiros, CRM ativa quem já levantou a mão. Na XP, a frente de ativação "
        "(8 meses, 2022) tirou da inércia contas abertas que ainda não movimentavam.",
    )
    grid(
        doc,
        ["Métrica", "Resultado"],
        [
            ["Ativação de contas", "+31,39%"],
            ["Receita estimada (jornadas e campanhas)", "+68,25%"],
        ],
        [9, 8],
    )
    callout(
        doc,
        "A PONTE COM A WORK BANK",
        "Na XP, a conta existia e o TED não saía. Na Work Bank, a indicação existe e o crédito não fecha. "
        "Mesma lógica: tirar da inércia quem já pediu. O CRM Vendas instala o funil para essa régua existir. "
        "Fonte: V4 Exclusive Experts · Frente de Ativação XP Inc. · 2022.",
    )

    add_heading(doc, "6. Como funciona na Work Bank")
    add_body(doc, "Um pipeline de crédito. Todo closer no mesmo jogo.")
    grid(
        doc,
        ["Etapa", "O que acontece"],
        [
            ["1. Indicação", "Entra no CRM com dono. Deixa de viver só no WhatsApp."],
            ["2. Simulação", "Produto, garantia e ticket no card. Closer assume."],
            ["3. Análise", "Parceiro e instituição. O histórico não se perde na troca."],
            ["4. Documentos", "Pendência visível. Follow-up com template, não de cabeça."],
            ["5. Desembolso", "O funil mostra o que fechou — e onde o crédito travou."],
        ],
        [4.2, 12.8],
    )

    add_heading(doc, "7. Escopo — CRM Vendas Basic")
    add_body(
        doc,
        "3 semanas · cerca de 13,5 horas. Esqueleto do processo comercial, não growth pleno. "
        "A plataforma de CRM já está definida internamente e não é nomeada neste documento.",
    )
    grid(
        doc,
        ["Entrega", "Detalhe"],
        [
            ["1º pipeline", "Estágios essenciais do crédito. Todo closer segue a mesma cadência."],
            ["Campos + time", "Mínimo viável para vendas, usuários e permissões do comercial."],
            ["3 templates", "E-mails SDR/Closer para o follow-up deixar de ser improvisado."],
            ["Importação da base", "Contatos existentes entram limpos e padronizados."],
            ["Treino", "Uso diário do funil e cadastro de oportunidades — prático, no sistema."],
            ["Documentação", "Configuração base repassada. O processo fica na ferramenta."],
        ],
        [4.5, 12.5],
    )

    add_heading(doc, "8. Rotina de 3 semanas")
    grid(
        doc,
        ["Semana", "Etapa", "Tarefa", "DRI"],
        [
            ["1", "Onboarding", "Grupo de comunicação + boas-vindas", "Account Manager"],
            ["1", "Onboarding", "Briefing inicial de vendas (processo e KPIs)", "Account Manager"],
            ["1", "Onboarding", "Reunião de alinhamento (escopo, cronograma, metodologia)", "Account Manager"],
            ["1", "Onboarding", "Acessos do time comercial e plataformas legadas", "Account Manager"],
            ["2", "Preparação", "Conta/ambiente no CRM de Vendas", "Profissional de CRM"],
            ["2", "Preparação", "1º pipeline de vendas (estágios essenciais)", "Profissional de CRM"],
            ["2", "Preparação", "Campos essenciais + usuários e permissões", "Profissional de CRM"],
            ["2", "Operação", "Importação da base (limpeza e padronização)", "Profissional de CRM"],
            ["2", "Operação", "3 templates de e-mail SDR/Closer", "CRM + Copy"],
            ["3", "Entrega", "Treinamento básico (funil e cadastro de oportunidades)", "Profissional de CRM"],
            ["3", "Entrega", "Documentação da configuração e repasse", "Profissional de CRM"],
            ["3", "Entrega", "Reunião de entrega e aprovação final", "Account Manager"],
        ],
        [2.2, 3.2, 8.2, 3.4],
    )

    add_heading(doc, "9. O que muda no dia 1")
    grid(
        doc,
        ["Quem", "O que muda"],
        [
            ["Closer", "Cadastra a oportunidade, move a etapa, usa o template. Follow-up deixa de depender de memória."],
            ["Gestor / sócio", "Vê o funil de crédito sem perguntar um a um. Sabe onde travou: simulação, parceiro ou doc."],
            ["Quem entra depois", "O processo está na ferramenta. Ramp-up deixa de ser 4 meses de sombra o sênior."],
        ],
        [4.2, 12.8],
    )
    callout(doc, "PROMESSA DE USO", "A metodologia vira rotina. O conhecimento não viaja com quem sai.")

    add_heading(doc, "10. Incluso")
    add_bullets(
        doc,
        [
            "Ambiente CRM Vendas: conta, usuários e permissões do time comercial",
            "1 pipeline de crédito com estágios essenciais mapeados",
            "Campos essenciais (mínimo viável para vendas)",
            "Importação da base com limpeza e padronização",
            "3 templates de e-mail SDR/Closer",
            "Treinamento básico + documentação e reunião de entrega",
        ],
    )

    add_heading(doc, "11. Por que agora e próximos passos")
    add_body(
        doc,
        "Setembro sem processo é o time novo herdando o artesanal. A janela é o onboarding comercial.",
    )
    grid(
        doc,
        ["Passo", "Ação"],
        [
            ["01 · Kickoff", "Grupo de comunicação e briefing de processo e KPIs de vendas."],
            ["02 · Acessos", "Time comercial, plataformas legadas e a base de contatos que já existe."],
            ["03 · 3 semanas", "Pipeline, templates, treino — e o funil rodando antes da rampa do time."],
        ],
        [4, 13],
    )

    add_heading(doc, "12. Investimento")
    price_block(doc)
    add_body(
        doc,
        "Documento elaborado pela V4 Company para a Work Bank. Valores conforme condição comercial "
        "acordada. A plataforma de CRM não é nomeada neste documento.",
        italic=True,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Work Bank × V4 Company  ·  Implementação CRM Vendas  ·  Confidencial")
    set_run_font(run, size=9, bold=True, color=RED)

    doc.save(OUT_DOCX)
    print("Wrote", OUT_DOCX)


def tr(cells, header=False):
    tds = []
    for i, c in enumerate(cells):
        style = 'font-weight:700;' if header or i == 0 else ""
        tds.append(f'<td style="{style}">{escape(c)}</td>')
    return "<tr>" + "".join(tds) + "</tr>"


def table_html(headers, rows):
    head = "".join(f"<th>{escape(h)}</th>" for h in headers)
    body = "".join(tr(r) for r in rows)
    return f'<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>'


def build_word_html():
    html = f"""<!DOCTYPE html>
<html xmlns:o="urn:schemas-microsoft-com:office:office"
      xmlns:w="urn:schemas-microsoft-com:office:word"
      xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<meta http-equiv="Content-Type" content="text/html; charset=utf-8">
<meta name="ProgId" content="Word.Document">
<meta name="Generator" content="Microsoft Word 15">
<title>Work Bank × V4 — Proposta CRM Vendas</title>
<!--[if gte mso 9]>
<xml>
  <w:WordDocument>
    <w:View>Print</w:View>
    <w:Zoom>100</w:Zoom>
    <w:DoNotOptimizeForBrowser/>
  </w:WordDocument>
</xml>
<![endif]-->
<style>
  @page WordSection1 {{ size: 21cm 29.7cm; margin: 1.8cm 2cm 2cm 2cm; }}
  div.WordSection1 {{ page: WordSection1; }}
  body {{
    font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #1A1A1A;
    line-height: 1.35; margin: 24px auto; max-width: 820px; background: #f4f0ef;
  }}
  div.WordSection1 {{ background: #fff; padding: 28px 32px 48px; }}
  h1 {{ font-size: 16pt; color: #A10F14; border-bottom: 1.5pt solid #A10F14; padding-bottom: 4pt; margin: 18pt 0 8pt; }}
  h2 {{ font-size: 12.5pt; color: #E50914; margin: 12pt 0 6pt; }}
  p {{ margin: 0 0 8pt; }}
  p.justify {{ text-align: justify; }}
  p.muted {{ color: #4A4A4A; font-size: 10.5pt; font-style: italic; }}
  p.footer-brand {{ color: #A10F14; font-weight: 700; text-align: center; font-size: 9pt; margin-top: 16pt; }}
  ul {{ margin: 4pt 0 10pt 18pt; padding: 0; }}
  li {{ margin: 0 0 3pt; color: #4A4A4A; }}
  table.grid {{ width: 100%; border-collapse: collapse; margin: 8pt 0 12pt; font-size: 10pt; }}
  table.grid th {{ background: #A10F14; color: #fff; text-align: left; font-size: 9pt; padding: 6pt 8pt; border: 1pt solid #A10F14; }}
  table.grid td {{ border: 1pt solid #E6D4D4; padding: 6pt 8pt; vertical-align: top; }}
  table.callout {{ width: 100%; margin: 8pt 0 12pt; border-left: 4pt solid #A10F14; background: #F8EFEF; }}
  table.callout td {{ padding: 10pt 12pt; }}
  .callout-label {{ color: #A10F14; font-size: 8.5pt; font-weight: 700; letter-spacing: .08em; margin: 0 0 4pt; }}
  table.cover {{ width: 100%; background: #280001; margin: 0 0 14pt; }}
  table.cover td {{ padding: 22pt 20pt; }}
  .cover-kicker {{ color: #F5B5B8; font-size: 9pt; font-weight: 700; letter-spacing: .12em; margin: 0 0 8pt; }}
  .cover-title {{ color: #fff; font-size: 26pt; font-weight: 700; margin: 0 0 6pt; }}
  .cover-sub {{ color: #FFD0D0; font-size: 14pt; margin: 0 0 10pt; }}
  .cover-lead {{ color: #E8D0D0; font-size: 11pt; margin: 0; }}
  table.price {{ width: 100%; margin: 8pt 0 12pt; }}
  td.price-left {{ background: #280001; color: #fff; padding: 16pt; vertical-align: top; width: 50%; }}
  td.price-right {{ background: #F8EFEF; border: 1.5pt solid #A10F14; padding: 14pt; vertical-align: top; width: 50%; }}
  .strike {{ color: #C9A0A0; text-decoration: line-through; font-size: 14pt; margin: 0 0 8pt; }}
  .price-now {{ color: #fff; font-size: 24pt; font-weight: 700; margin: 0 0 8pt; }}
</style>
</head>
<body>
<div class="WordSection1">
<table class="cover" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="cover-kicker">PROPOSTA COMERCIAL</p>
<p class="cover-title">CRM Vendas</p>
<p class="cover-sub">Work Bank  ×  V4 Company</p>
<p class="cover-lead">Fundação de CRM Vendas em 3 semanas — pipeline de crédito, templates e rotina para o time que entra em setembro.</p>
</td></tr></table>
{table_html(["Campo", "Informação"], [
    ["Cliente", "Work Bank (Marchi Soluções Financeiras)"],
    ["Produto", "Implementação CRM Vendas Basic"],
    ["Prazo", "3 semanas  ·  ~13,5 horas"],
    ["Investimento", "R$ 5.639,21  (tabela: R$ 9.145,16)"],
    ["Condução", "V4 Company  ·  Account Manager + Profissional de CRM + Copy (templates)"],
    ["Momento", "Time comercial novo entra em setembro"],
])}
<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">LEITURA CENTRAL</p>
<p>A Work Bank não tem problema de demanda. Tem um comercial artesanal: 90% indicação, funil na cabeça do closer e time novo sem processo na ferramenta. Esta proposta instala o esqueleto do CRM de Vendas — um pipeline de crédito, templates e rotina — em 3 semanas.</p>
</td></tr></table>

<h1>1. Objetivo</h1>
<p class="justify">Instalar a fundação comercial da Work Bank no CRM de Vendas: um pipeline mapeado no ciclo de crédito, campos essenciais, usuários do time, importação da base, 3 templates de e-mail SDR/Closer, treinamento básico e documentação da configuração. O resultado esperado não é um manual na gaveta. É o closer cadastrando oportunidade e o gestor enxergando o funil no dia 1.</p>

<h1>2. O momento Work Bank</h1>
<p class="justify">Assessoria de crédito (B2B e B2C), mais de 80 instituições, atuação em Minas Gerais e entorno. Produtos: home equity, garantia de veículo, capital de giro e consórcios. Receita atual de cerca de R$ 80 mil/mês, com ambição de R$ 20 milhões em desembolso.</p>
{table_html(["Dado", "Leitura"], [
    ["R$ 80 mil/mês", "Modelo validado. A máquina ainda é artesanal."],
    ["90% indicação", "Networking segura a operação. Sem funil, a previsão some."],
    ["Setembro", "Time comercial novo entra. Sem processo no CRM, cada um inventa o próprio."],
    ["R$ 20 milhões", "Ambição de desembolso. O closer precisa de pipeline no dia 1."],
])}

<h1>3. Diagnóstico</h1>
<p class="justify">Três vazamentos travam a previsibilidade. Começar errado — CRM como bloco de notas — faz o time de setembro herdar o vício.</p>
{table_html(["Vazamento", "Efeito"], [
    ["Oportunidade sem dono", "Indicação entra no WhatsApp. Ninguém registra. Follow-up depende de memória."],
    ["Funil na cabeça", "Cada closer vende de um jeito. Gestor não confia no número. Win rate vira opinião."],
    ["Time novo sem rampa", "Sem processo no sistema, onboarding pode ir a 4 meses. O conhecimento sai com a pessoa."],
])}
<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">O RISCO</p>
<p>A indicação existe e o crédito não fecha. Sem o 1º pipeline, o gestor não sabe por que o crédito morreu: simulação, parceiro ou documentação.</p>
</td></tr></table>

<h1>4. O que é o CRM de Vendas</h1>
<p class="justify">Não é agenda de contato. É o esqueleto do processo comercial: o lugar onde cada pedido de crédito vive — do primeiro “me indica” até o desembolso. A metodologia sai da cabeça do vendedor e vira rotina na ferramenta.</p>
{table_html(["Peça", "Para que serve"], [
    ["Dono", "Toda oportunidade tem responsável."],
    ["Etapa", "O crédito anda no funil, não no chat."],
    ["Histórico", "O gestor vê o funil, não pergunta um a um."],
    ["Rotina", "O closer novo herda o processo no dia 1."],
])}

<h1>5. Case — XP Inc. × V4 Exclusive Experts</h1>
<p class="justify">Em serviços financeiros, CRM ativa quem já levantou a mão. Na XP, a frente de ativação (8 meses, 2022) tirou da inércia contas abertas que ainda não movimentavam.</p>
{table_html(["Métrica", "Resultado"], [
    ["Ativação de contas", "+31,39%"],
    ["Receita estimada (jornadas e campanhas)", "+68,25%"],
])}
<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">A PONTE COM A WORK BANK</p>
<p>Na XP, a conta existia e o TED não saía. Na Work Bank, a indicação existe e o crédito não fecha. Mesma lógica: tirar da inércia quem já pediu. O CRM Vendas instala o funil para essa régua existir. Fonte: V4 Exclusive Experts · Frente de Ativação XP Inc. · 2022.</p>
</td></tr></table>

<h1>6. Como funciona na Work Bank</h1>
<p class="justify">Um pipeline de crédito. Todo closer no mesmo jogo.</p>
{table_html(["Etapa", "O que acontece"], [
    ["1. Indicação", "Entra no CRM com dono. Deixa de viver só no WhatsApp."],
    ["2. Simulação", "Produto, garantia e ticket no card. Closer assume."],
    ["3. Análise", "Parceiro e instituição. O histórico não se perde na troca."],
    ["4. Documentos", "Pendência visível. Follow-up com template, não de cabeça."],
    ["5. Desembolso", "O funil mostra o que fechou — e onde o crédito travou."],
])}

<h1>7. Escopo — CRM Vendas Basic</h1>
<p class="justify">3 semanas · cerca de 13,5 horas. Esqueleto do processo comercial, não growth pleno. A plataforma de CRM já está definida internamente e não é nomeada neste documento.</p>
{table_html(["Entrega", "Detalhe"], [
    ["1º pipeline", "Estágios essenciais do crédito. Todo closer segue a mesma cadência."],
    ["Campos + time", "Mínimo viável para vendas, usuários e permissões do comercial."],
    ["3 templates", "E-mails SDR/Closer para o follow-up deixar de ser improvisado."],
    ["Importação da base", "Contatos existentes entram limpos e padronizados."],
    ["Treino", "Uso diário do funil e cadastro de oportunidades — prático, no sistema."],
    ["Documentação", "Configuração base repassada. O processo fica na ferramenta."],
])}

<h1>8. Rotina de 3 semanas</h1>
{table_html(["Semana", "Etapa", "Tarefa", "DRI"], [
    ["1", "Onboarding", "Grupo de comunicação + boas-vindas", "Account Manager"],
    ["1", "Onboarding", "Briefing inicial de vendas (processo e KPIs)", "Account Manager"],
    ["1", "Onboarding", "Reunião de alinhamento (escopo, cronograma, metodologia)", "Account Manager"],
    ["1", "Onboarding", "Acessos do time comercial e plataformas legadas", "Account Manager"],
    ["2", "Preparação", "Conta/ambiente no CRM de Vendas", "Profissional de CRM"],
    ["2", "Preparação", "1º pipeline de vendas (estágios essenciais)", "Profissional de CRM"],
    ["2", "Preparação", "Campos essenciais + usuários e permissões", "Profissional de CRM"],
    ["2", "Operação", "Importação da base (limpeza e padronização)", "Profissional de CRM"],
    ["2", "Operação", "3 templates de e-mail SDR/Closer", "CRM + Copy"],
    ["3", "Entrega", "Treinamento básico (funil e cadastro de oportunidades)", "Profissional de CRM"],
    ["3", "Entrega", "Documentação da configuração e repasse", "Profissional de CRM"],
    ["3", "Entrega", "Reunião de entrega e aprovação final", "Account Manager"],
])}

<h1>9. O que muda no dia 1</h1>
{table_html(["Quem", "O que muda"], [
    ["Closer", "Cadastra a oportunidade, move a etapa, usa o template. Follow-up deixa de depender de memória."],
    ["Gestor / sócio", "Vê o funil de crédito sem perguntar um a um. Sabe onde travou: simulação, parceiro ou doc."],
    ["Quem entra depois", "O processo está na ferramenta. Ramp-up deixa de ser 4 meses de sombra o sênior."],
])}
<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">PROMESSA DE USO</p>
<p>A metodologia vira rotina. O conhecimento não viaja com quem sai.</p>
</td></tr></table>

<h1>10. Incluso</h1>
<ul>
<li>Ambiente CRM Vendas: conta, usuários e permissões do time comercial</li>
<li>1 pipeline de crédito com estágios essenciais mapeados</li>
<li>Campos essenciais (mínimo viável para vendas)</li>
<li>Importação da base com limpeza e padronização</li>
<li>3 templates de e-mail SDR/Closer</li>
<li>Treinamento básico + documentação e reunião de entrega</li>
</ul>

<h1>11. Por que agora e próximos passos</h1>
<p class="justify">Setembro sem processo é o time novo herdando o artesanal. A janela é o onboarding comercial.</p>
{table_html(["Passo", "Ação"], [
    ["01 · Kickoff", "Grupo de comunicação e briefing de processo e KPIs de vendas."],
    ["02 · Acessos", "Time comercial, plataformas legadas e a base de contatos que já existe."],
    ["03 · 3 semanas", "Pipeline, templates, treino — e o funil rodando antes da rampa do time."],
])}

<h1>12. Investimento</h1>
<table class="price" border="0" cellspacing="0" cellpadding="0">
<tr>
<td class="price-left" width="50%">
<p class="cover-kicker">TABELA · IMPLEMENTAÇÃO CRM VENDAS</p>
<p class="strike">R$ 9.145,16</p>
<p class="cover-kicker">CONDIÇÃO WORK BANK</p>
<p class="price-now">R$ 5.639,21</p>
<p class="cover-lead">3 semanas · CRM Vendas Basic<br>pipeline, templates, treino e documentação</p>
</td>
<td class="price-right" width="50%">
<p class="callout-label">O QUE ESTÁ INCLUSO</p>
<p><strong>Funil instalado.</strong> Pipeline de crédito, campos, base, 3 templates e o time treinado no uso diário.</p>
<p><strong>Por que agora.</strong> O closer novo herda processo na ferramenta, não no WhatsApp.</p>
<p><strong>Próximo passo.</strong> Briefing + acessos. Em 3 semanas o funil está no ar.</p>
</td>
</tr>
</table>
<p class="muted">Documento elaborado pela V4 Company para a Work Bank. Valores conforme condição comercial acordada. A plataforma de CRM não é nomeada neste documento.</p>
<p class="footer-brand">Work Bank × V4 Company  ·  Implementação CRM Vendas  ·  Confidencial</p>
</div>
</body>
</html>
"""
    OUT_HTML.write_text(html, encoding="utf-8")
    copy2(OUT_HTML, OUT_DOC)
    print("Wrote", OUT_HTML)
    print("Wrote", OUT_DOC)


if __name__ == "__main__":
    build_docx()
    build_word_html()
