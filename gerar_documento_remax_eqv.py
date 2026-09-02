#!/usr/bin/env python3
"""Gera a proposta EQV da Remax Pro Work em HTML Word + .doc (abre no Word) + .docx."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CLIENT = "Remax Pro Work"
TABELA = "R$ 50.632,85"
AVISTA = "R$ 19.500,00"
PARCELA = "R$ 3.500,00"
PARCELAS = "6x de R$ 3.500,00"
CASHBACK = "R$ 15.000,00"
TOTAL_PARCELADO = "R$ 21.000,00"
ECONOMIA = "R$ 31.132,85"

RED = "A10F14"
RED2 = "E50914"
INK = "1A1A1A"
MUT = "4A4A4A"
SOFT = "F8EFEF"
LINE = "E6D4D4"
WHITE = "FFFFFF"
DARK = "280001"

OUT_DIR = Path("/workspace/remax-ok")
OUT_HTML = OUT_DIR / "documento.html"
OUT_DOC = OUT_DIR / "Remax-Pro-Work-EQV.doc"
OUT_DOCX = OUT_DIR / "Remax-Pro-Work-EQV.docx"
OUT_LANDING = OUT_DIR / "word.html"

BRANCH = "cursor/remax-pro-work-eqv-061f"
RAW_DOC = f"https://github.com/sarapizzico-hue/Sara/raw/{BRANCH}/remax-ok/Remax-Pro-Work-EQV.doc"
RAW_DOCX = f"https://github.com/sarapizzico-hue/Sara/raw/{BRANCH}/remax-ok/Remax-Pro-Work-EQV.docx"
PREVIEW = (
    "https://htmlpreview.github.io/?https://github.com/sarapizzico-hue/Sara/blob/"
    f"{BRANCH}/remax-ok/documento.html"
)


def html_doc() -> str:
    return f"""<!DOCTYPE html>
<html xmlns:o="urn:schemas-microsoft-com:office:office"
      xmlns:w="urn:schemas-microsoft-com:office:word"
      xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<meta http-equiv="Content-Type" content="text/html; charset=utf-8">
<meta name="ProgId" content="Word.Document">
<meta name="Generator" content="Microsoft Word 15">
<title>{CLIENT} × V4 — Empresa que Vende</title>
<!--[if gte mso 9]>
<xml>
  <w:WordDocument>
    <w:View>Print</w:View>
    <w:Zoom>100</w:Zoom>
    <w:DoNotOptimizeForBrowser/>
  </w:WordDocument>
</xml>
<![endif]-->
<style>
  @page WordSection1 {{ size: 21cm 29.7cm; margin: 1.8cm 2cm 2cm 2cm; }}
  div.WordSection1 {{ page: WordSection1; }}
  body {{
    font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #1A1A1A;
    line-height: 1.35; margin: 24px auto; max-width: 820px; background: #f4f0ef;
  }}
  div.WordSection1 {{ background: #fff; padding: 28px 32px 48px; }}
  h1 {{ font-size: 16pt; color: #A10F14; border-bottom: 1.5pt solid #A10F14; padding-bottom: 4pt; margin: 18pt 0 8pt; }}
  h2 {{ font-size: 12.5pt; color: #E50914; margin: 12pt 0 6pt; }}
  p {{ margin: 0 0 8pt; }}
  p.justify {{ text-align: justify; }}
  p.muted {{ color: #4A4A4A; font-size: 10.5pt; font-style: italic; }}
  p.label {{ color: #E50914; font-weight: 700; margin: 8pt 0 4pt; }}
  p.footer-brand {{ color: #A10F14; font-weight: 700; text-align: center; font-size: 9pt; margin-top: 16pt; }}
  ul {{ margin: 4pt 0 10pt 18pt; padding: 0; }}
  li {{ margin: 0 0 3pt; color: #4A4A4A; }}
  table.grid {{ width: 100%; border-collapse: collapse; margin: 8pt 0 12pt; font-size: 10pt; }}
  table.grid th {{ background: #A10F14; color: #fff; text-align: left; font-size: 9pt; padding: 6pt 8pt; border: 1pt solid #A10F14; }}
  table.grid td {{ border: 1pt solid #E6D4D4; padding: 6pt 8pt; vertical-align: top; }}
  table.callout {{ width: 100%; margin: 8pt 0 12pt; border-left: 4pt solid #A10F14; background: #F8EFEF; }}
  table.callout td {{ padding: 10pt 12pt; }}
  .callout-label {{ color: #A10F14; font-size: 8.5pt; font-weight: 700; letter-spacing: .08em; margin: 0 0 4pt; }}
  table.cover {{ width: 100%; background: #280001; margin: 0 0 14pt; }}
  table.cover td {{ padding: 22pt 20pt; }}
  .cover-kicker {{ color: #F5B5B8; font-size: 9pt; font-weight: 700; letter-spacing: .12em; margin: 0 0 8pt; }}
  .cover-title {{ color: #fff; font-size: 26pt; font-weight: 700; margin: 0 0 6pt; }}
  .cover-sub {{ color: #FFD0D0; font-size: 14pt; margin: 0 0 10pt; }}
  .cover-lead {{ color: #E8D0D0; font-size: 11pt; margin: 0; }}
  table.price {{ width: 100%; margin: 8pt 0 12pt; }}
  td.price-left {{ background: #280001; color: #fff; padding: 16pt; vertical-align: top; width: 50%; }}
  td.price-right {{ background: #F8EFEF; border: 1.5pt solid #A10F14; padding: 14pt; vertical-align: top; width: 50%; }}
  .strike {{ color: #C9A0A0; text-decoration: line-through; font-size: 14pt; margin: 0 0 8pt; }}
  .price-now {{ color: #fff; font-size: 24pt; font-weight: 700; margin: 0 0 8pt; }}
  .price-month {{ font-size: 13pt; font-weight: 700; margin: 0 0 8pt; }}
</style>
</head>
<body>
<div class="WordSection1">
<table class="cover" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="cover-kicker">PROPOSTA COMERCIAL</p>
<p class="cover-title">Empresa que Vende</p>
<p class="cover-sub">{CLIENT}  ×  V4 Company</p>
<p class="cover-lead">Consultoria, acompanhamento e execução assistida para transformar marketing e comercial em uma máquina de receita previsível — em 6 meses.</p>
</td></tr></table>

<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Campo</th><th>Informação</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Cliente</td><td>{CLIENT}</td></tr>
<tr><td style="font-weight:700;">Produto</td><td>Empresa que Vende (EQV) — programa semestral de consultoria + execução assistida</td></tr>
<tr><td style="font-weight:700;">Prazo</td><td>6 meses  ·  5 fases  ·  squad de 4 especialistas (Conteúdo, Mídia Paga, Comercial, CRM)</td></tr>
<tr><td style="font-weight:700;">Investimento</td><td>{PARCELAS}  ·  PIX à vista {AVISTA}  ·  referência {TABELA}</td></tr>
<tr><td style="font-weight:700;">Cashback EE</td><td>{CASHBACK} já investidos na Estruturação Estratégica, aplicados nesta condição</td></tr>
<tr><td style="font-weight:700;">Condução</td><td>V4 Company  ·  consultor estratégico (1h/mês) + squad de execução assistida</td></tr>
<tr><td style="font-weight:700;">Incluso</td><td>Comunidade, monitoria coletiva mensal, conteúdo semanal e Fábrica de Receita (imersão presencial anual)</td></tr>
<tr><td style="font-weight:700;">Garantia</td><td>30 dias de serviço — reembolso incondicional se a entrega não fizer sentido</td></tr>
</tbody></table>

<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">LEITURA CENTRAL</p>
<p>Você trabalha muito. A receita ainda não é previsível. O Empresa que Vende não executa no seu lugar: estrutura o negócio, treina o time e acompanha a execução. Quem opera é você. É exatamente por isso que a máquina fica na empresa quando o ciclo acaba.</p>
</td></tr></table>

<h1>1. Objetivo</h1>
<p class="justify">Instalar, em 6 meses, o método para a {CLIENT} vender com previsibilidade: diagnóstico contínuo de maturidade GTM, consultoria estratégica mensal, execução assistida nas quatro frentes (conteúdo, mídia, comercial e CRM), comunidade de pares e imersão presencial da Fábrica de Receita.</p>
<p class="justify">O resultado esperado não é um PDF para a gaveta nem uma agência que some quando o contrato acaba. É o time operando com tese, artefato e cadência — conteúdo que gera demanda, mídia decidida por dado, comercial que fecha e CRM que governa o pipeline.</p>

<h1>2. O problema que o programa resolve</h1>
<p class="justify">O empresário típico do EQV já vende. O que falta é sistema para repetir o mês. Vendas oscilam; o processo mora na cabeça do melhor vendedor; marketing decide no achismo; a agência executa e o crescimento para quando o contrato termina.</p>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Sintoma</th><th>Efeito na operação</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Receita imprevisível</td><td>O mês é esforço, não método. Sem processo, o resultado não se repete.</td></tr>
<tr><td style="font-weight:700;">Processo na cabeça de uma pessoa</td><td>Quando o melhor vendedor sai, o comercial sai junto.</td></tr>
<tr><td style="font-weight:700;">Marketing no achismo</td><td>Canal, conteúdo e mídia sem tese. Investe, testa, para.</td></tr>
<tr><td style="font-weight:700;">Agência executa — você não aprende</td><td>Quando o contrato acaba, o resultado vai junto. A operação nunca herdou a máquina.</td></tr>
<tr><td style="font-weight:700;">Comercial amador</td><td>Script, follow-up e pipeline improvisados. O funil vaza depois do lead.</td></tr>
<tr><td style="font-weight:700;">Posicionamento frouxo</td><td>A marca não escolheu ICP. O mercado escolhe preço.</td></tr>
</tbody></table>

<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">POR QUE AS SOLUÇÕES TRADICIONAIS NÃO RESOLVEM</p>
<p><strong>Agência</strong> executa por você — o resultado some com o contrato. <strong>Curso</strong> ensina teoria sem aplicar no negócio real. <strong>Consultoria clássica</strong> diagnostica e entrega relatório; ninguém implementa. O empresário continua sem máquina de vendas previsível.</p>
</td></tr></table>

<h1>3. Para quem é o Empresa que Vende</h1>
<p class="justify">O programa foi desenhado para operação Tiny/Small que já tem venda ativa e ainda não tem sistema. Não é “começar a vender”. É estruturar o que já existe — antes de um fee de execução cheio.</p>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Filtro</th><th>Perfil</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Porte</td><td>Tiny e Small · até 50 pessoas</td></tr>
<tr><td style="font-weight:700;">Faturamento</td><td>R$ 600 mil a R$ 2,4 milhões ao ano</td></tr>
<tr><td style="font-weight:700;">Mercado</td><td>B2B / serviços com operação de vendas ativa</td></tr>
<tr><td style="font-weight:700;">Canal</td><td>Inside sales · WhatsApp · Meta Ads · CRM</td></tr>
<tr><td style="font-weight:700;">Maturidade GTM</td><td>Baixo a médio (N1 ou N2 no diagnóstico V4)</td></tr>
</tbody></table>
<p class="label">Sinais de encaixe</p>
<ul>
<li>Experiências frustrantes com agências — pagou para executar e não herdou método.</li>
<li>Setor de marketing pouco experiente, sem tese nem leitura de dado.</li>
<li>Estrutura comercial amadora: script, follow-up e pipeline soltos.</li>
<li>Sem plano de marketing definido; prioridade muda no feeling.</li>
<li>Sem posicionamento claro; compete por preço.</li>
</ul>

<h1>4. A solução</h1>
<p class="justify">A ordem de construção é clareza → sistema → execução. Sem método, mídia vira aposta e comercial vira pressão. Com método, receita vira consequência.</p>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Pilar</th><th>O que instala</th></tr></thead><tbody>
<tr><td style="font-weight:700;">01 · Diagnóstico contínuo</td><td>Maturidade GTM monitorada. Gargalos identificados a cada ciclo — sem achismo.</td></tr>
<tr><td style="font-weight:700;">02 · Consultoria estratégica</td><td>1h/mês com consultor V4: define prioridades e aloca os créditos do squad.</td></tr>
<tr><td style="font-weight:700;">03 · Execução assistida</td><td>Squad de 4 especialistas: Conteúdo, Mídia Paga, Comercial e CRM.</td></tr>
<tr><td style="font-weight:700;">04 · Comunidade</td><td>Peers do mesmo estágio, lives mensais temáticas e grupo fechado no WhatsApp.</td></tr>
<tr><td style="font-weight:700;">05 · Fábrica de Receita</td><td>Imersão presencial anual: hot seats, cases reais e workshops com a V4.</td></tr>
</tbody></table>
<p class="justify">Papéis claros: a V4 prescreve e conduz; o cliente opera no dia a dia. O consultor aloca as horas do squad no que destrava a receita naquele ciclo — não empurra pacote.</p>

<h1>5. Jornada de 6 meses</h1>
<p class="justify">Cinco fases. Um ciclo contínuo. A máquina é instalada no time, não alugada.</p>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Fase</th><th>Quando</th><th>Nome</th><th>Função</th></tr></thead><tbody>
<tr><td style="font-weight:700;">01</td><td>Mês 1</td><td>Onboarding</td><td>Kickoff, maturidade GTM e definição de prioridades do ciclo.</td></tr>
<tr><td style="font-weight:700;">02</td><td>Mês 2</td><td>EE 3.0 completa</td><td>3P-IA, posicionamento ativo e canal de vendas no ar.</td></tr>
<tr><td style="font-weight:700;">03</td><td>Mês 3</td><td>Consultoria de conteúdo</td><td>Produção assistida + linha editorial no método.</td></tr>
<tr><td style="font-weight:700;">04</td><td>Mês 4</td><td>Consultoria de tráfego</td><td>Campanhas, pixel, audiência e otimização por resultado.</td></tr>
<tr><td style="font-weight:700;">05</td><td>Mês 5/6</td><td>Consultoria comercial + CRM</td><td>Scripts, follow-up, pipeline e monetização do funil completo.</td></tr>
</tbody></table>

<h2>5.1 Sessão 01 — Produto, oferta, posicionamento e conteúdo</h2>
<p class="justify">O especialista V4 orienta a estratégia, calibra a mensagem para o ICP e acompanha a execução mês a mês. A marca para de postar no feeling e passa a gerar demanda com tese.</p>
<p class="label">Artefatos</p>
<ul>
<li>Roteiros de vídeo viral calibrados (prontos para gravar).</li>
<li>Brief de posts complementares para Instagram.</li>
<li>Modelo de linha editorial.</li>
<li>Funil de creator-led growth.</li>
</ul>

<h2>5.2 Sessão 02 — Mídia e performance</h2>
<p class="justify">Acompanhamento para operar mídia paga com método: escolha de canais, leitura de dados e decisão de investimento. Você para de apostar em canal e passa a decidir por ROAS e CPL.</p>
<p class="label">Artefatos</p>
<ul>
<li>Configuração e leitura de pixel.</li>
<li>Plano de mídia estruturado e segmentação de audiências.</li>
<li>Interpretação de dados das campanhas.</li>
<li>Capacitação para gestão de mídia paga e tomada de decisão por resultado.</li>
</ul>

<h2>5.3 Sessão 03 — Comercial e CRM</h2>
<p class="justify">O especialista revisa scripts, analisa as objeções reais do mês e ajusta o processo de vendas com o time. O comercial deixa de ser artesanal. Fecha com cadência.</p>
<p class="label">Artefatos</p>
<ul>
<li>Revisão dos scripts de venda e cliente oculto.</li>
<li>Análise de objeções reais do mês e treinamento comercial.</li>
<li>Atualização das réguas de conversão e follow-up.</li>
<li>CRM ativo: cadência, pipeline e conversão — não agenda.</li>
</ul>

<h1>6. O que o squad instala</h1>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Frente</th><th>Entrega</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Conteúdo</td><td>Roteiro de vídeo viral calibrado ao negócio, briefing de posts Instagram, suporte à produção e publicação.</td></tr>
<tr><td style="font-weight:700;">Mídia paga</td><td>Campanhas Meta/Google, segmentação, pixel, leitura de dados, decisão por ROAS e CPL.</td></tr>
<tr><td style="font-weight:700;">Comercial</td><td>Scripts, análise de objeções, role play com vendedores, réguas de follow-up.</td></tr>
<tr><td style="font-weight:700;">CRM</td><td>CRM ativo, fluxos de cadência e automações, análise de pipeline e conversão.</td></tr>
</tbody></table>

<h1>7. Comunidade e Fábrica de Receita</h1>
<p class="justify">Além das sessões de execução assistida, o programa inclui acesso permanente a:</p>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Camada</th><th>O que entra</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Monitoria coletiva mensal</td><td>Encontros com especialistas V4 e peers para ajustar rota juntos.</td></tr>
<tr><td style="font-weight:700;">Comunidade WhatsApp</td><td>Grupo fechado para troca rápida, alinhamentos e suporte contínuo.</td></tr>
<tr><td style="font-weight:700;">Conteúdo semanal</td><td>Aquisição, retenção e monetização para manter o time no método.</td></tr>
<tr><td style="font-weight:700;">Fábrica de Receita</td><td>2 dias de imersão presencial anual: hot seats, workshops e networking de alto nível.</td></tr>
</tbody></table>

<h1>8. Por que a V4</h1>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>#</th><th>Princípio</th></tr></thead><tbody>
<tr><td style="font-weight:700;">1</td><td>Você para de apostar e começa a operar com método — menos tentativa, mais decisão testada.</td></tr>
<tr><td style="font-weight:700;">2</td><td>Operamos com dados de mercado + experiência real: milhares de operações e sistemas de receita, não achismo.</td></tr>
<tr><td style="font-weight:700;">3</td><td>Não empurramos serviços. Prescrevemos o necessário para destravar a receita agora.</td></tr>
<tr><td style="font-weight:700;">4</td><td>Garantia incondicional de risco zero: se o método não fizer sentido, o risco é nosso.</td></tr>
</tbody></table>
<p class="muted">Como já fizemos com Ortobom, Spotify, GioLaser, Lugano, XP Inc. e Smart Fit — entre outros.</p>

<h1>9. Investimento</h1>
<table class="price" border="0" cellspacing="0" cellpadding="0">
<tr>
<td class="price-left" width="50%">
<p class="cover-kicker">VALOR DE REFERÊNCIA</p>
<p class="strike">{TABELA}</p>
<p class="cover-kicker">PIX / À VISTA</p>
<p class="price-now">{AVISTA}</p>
<p class="cover-lead">Programa semestral · consultoria + squad<br>comunidade + Fábrica de Receita</p>
</td>
<td class="price-right" width="50%">
<p class="callout-label">CONDIÇÕES</p>
<p class="price-month">{PARCELAS} /mês</p>
<p>PIX à vista {AVISTA}. Parcelado sem juros: {PARCELAS} (total {TOTAL_PARCELADO}).</p>
<p>Cashback de {CASHBACK} da Estruturação Estratégica já aplicado nesta condição.</p>
<p>Custo do ciclo de 6 meses. A máquina fica no time. Não é aluguel de agência.</p>
</td>
</tr>
</table>

<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Linha</th><th>Valor</th></tr></thead><tbody>
<tr><td style="font-weight:700;">Valor de tabela / referência</td><td>{TABELA}</td></tr>
<tr><td style="font-weight:700;">PIX / à vista</td><td>{AVISTA}</td></tr>
<tr><td style="font-weight:700;">Parcelado</td><td>{PARCELAS}  ·  total {TOTAL_PARCELADO}</td></tr>
<tr><td style="font-weight:700;">Cashback EE aplicado</td><td>{CASHBACK}</td></tr>
<tr><td style="font-weight:700;">Economia no PIX vs. tabela</td><td>{ECONOMIA}</td></tr>
</tbody></table>

<table class="callout" border="0" cellspacing="0" cellpadding="0"><tr><td>
<p class="callout-label">LEITURA DE DECISÃO</p>
<p>O ciclo instala conteúdo, mídia, comercial e CRM no mesmo programa — com o time herdando o método. O valor já investido na EE entra como cashback. Adiar o EQV mantém o achismo: mídia continua tentando, comercial continua artesanal, o dono continua sendo o processo.</p>
</td></tr></table>

<h2>9.1 Garantia</h2>
<p class="justify">Se após 30 dias de serviço você não estiver satisfeito com a entrega, devolvemos o investimento efetuado. Risco zero para entrar. O método precisa fazer sentido na operação — não só no papel.</p>

<h1>10. Próximos passos</h1>
<table class="grid" border="1" cellspacing="0" cellpadding="0"><thead><tr><th>Responsável</th><th>Ação</th></tr></thead><tbody>
<tr><td style="font-weight:700;">{CLIENT}</td><td>Validar a condição comercial (PIX {AVISTA} ou {PARCELAS}) e formalizar o start.</td></tr>
<tr><td style="font-weight:700;">V4</td><td>Kickoff do mês 01: diagnóstico de maturidade GTM, prioridades do ciclo e alocação do squad.</td></tr>
<tr><td style="font-weight:700;">V4 + {CLIENT}</td><td>Entrar na comunidade, na monitoria mensal e no calendário da Fábrica de Receita a partir da assinatura.</td></tr>
</tbody></table>

<p class="muted">Documento elaborado pela V4 Company para a {CLIENT}. Valores conforme condição comercial da base ativa, com cashback de {CASHBACK} da Estruturação Estratégica aplicado sobre a referência de {TABELA}.</p>
<p class="footer-brand">{CLIENT} × V4 Company  ·  Empresa que Vende  ·  Confidencial</p>
</div>
</body>
</html>
"""


def landing_html() -> str:
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{CLIENT} · Baixar Word | Empresa que Vende</title>
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link href="https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;500;600;700&family=IBM+Plex+Mono:wght@500;600&display=swap" rel="stylesheet" />
  <style>
    :root {{ --red:#e50914; --dark:#280001; --ink:#1a1a1a; --mut:#5a5a5a; --paper:#f7f5f4; }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0; min-height: 100vh;
      font: 16px/1.5 "IBM Plex Sans", system-ui, sans-serif;
      color: var(--ink);
      background:
        radial-gradient(ellipse 70% 40% at 50% -10%, rgba(229,9,20,.18), transparent 55%),
        linear-gradient(180deg, #fff 0%, var(--paper) 100%);
      display: flex; align-items: center; justify-content: center;
      padding: 28px 16px;
    }}
    .card {{
      width: 100%; max-width: 640px;
      background: #fff;
      border: 1px solid #ead4d4;
      border-radius: 24px;
      padding: 36px 32px 28px;
      box-shadow: 0 24px 70px rgba(40,0,1,.08);
    }}
    .kicker {{
      font: 600 11px/1 "IBM Plex Mono", monospace;
      letter-spacing: .12em; text-transform: uppercase;
      color: var(--red); margin: 0 0 10px;
    }}
    h1 {{ margin: 0 0 8px; font-size: clamp(26px, 5vw, 34px); letter-spacing: -.03em; }}
    .lead {{ margin: 0 0 22px; color: var(--mut); }}
    .btns {{ display: grid; gap: 10px; margin-bottom: 18px; }}
    a.btn {{
      display: block; text-align: center; text-decoration: none;
      border-radius: 14px; padding: 16px 18px; font-weight: 700;
      transition: transform .15s ease, box-shadow .15s ease;
    }}
    a.btn:hover {{ transform: translateY(-1px); }}
    .btn-word {{
      background: var(--dark); color: #fff;
      box-shadow: 0 12px 28px rgba(40,0,1,.22);
    }}
    .btn-google {{
      background: #fff; color: var(--dark);
      border: 2px solid var(--red);
    }}
    .btn-google span {{ color: var(--red); }}
    .btn-docx {{
      background: #f8efef; color: var(--ink); border: 1px solid #ead4d4;
      font-weight: 600;
    }}
    .share {{
      background: #faf7f7; border: 1px dashed #d7b4b4;
      border-radius: 14px; padding: 12px 14px; margin-top: 6px;
    }}
    .share label {{
      display: block; font: 600 10px/1 "IBM Plex Mono", monospace;
      letter-spacing: .1em; text-transform: uppercase; color: var(--red);
      margin-bottom: 8px;
    }}
    .share-row {{ display: flex; gap: 8px; }}
    .share input {{
      flex: 1; border: 1px solid #ead4d4; border-radius: 10px;
      padding: 10px 12px; font-size: 13px; color: var(--ink);
      background: #fff;
    }}
    .share button {{
      border: 0; border-radius: 10px; background: var(--red); color: #fff;
      font-weight: 700; padding: 10px 14px; cursor: pointer;
    }}
    ol {{ margin: 14px 0 0; padding-left: 18px; color: var(--mut); font-size: 14px; }}
    li {{ margin: 4px 0; }}
    .ok {{ display: none; margin: 8px 0 0; color: #167846; font-weight: 700; font-size: 13px; }}
    .ok.show {{ display: block; }}
  </style>
</head>
<body>
  <div class="card">
    <p class="kicker">{CLIENT} × V4 Company</p>
    <h1>Empresa que Vende</h1>
    <p class="lead">Proposta comercial do EQV. O .doc abre direto no Microsoft Word. Também dá para subir no Google Docs.</p>

    <div class="btns">
      <a class="btn btn-word" id="btn-word" href="{RAW_DOC}">
        Baixar .doc (abre direto no Word)
      </a>
      <a class="btn btn-google" id="btn-google" target="_blank" rel="noopener"
         href="https://docs.google.com/gview?url={RAW_DOCX}&amp;embedded=false">
        Subir no <span>Google Docs</span>
      </a>
      <a class="btn btn-docx" href="{RAW_DOCX}">
        Baixar .docx
      </a>
    </div>

    <div class="share">
      <label>Link para enviar</label>
      <div class="share-row">
        <input id="share-url" readonly value="https://htmlpreview.github.io/?https://github.com/sarapizzico-hue/Sara/blob/{BRANCH}/remax-ok/word.html" />
        <button type="button" id="copy">Copiar</button>
      </div>
      <p class="ok" id="copied">Link copiado.</p>
    </div>

    <ol>
      <li><strong>Baixar .doc</strong> — baixa e abre direto no Microsoft Word.</li>
      <li><strong>Subir no Google Docs</strong> — abre no visualizador. Clique em <em>Abrir com o Documentos Google</em>.</li>
      <li>Valores: {PARCELAS} · PIX {AVISTA} · cashback EE {CASHBACK}.</li>
    </ol>
  </div>

  <script>
    (function () {{
      var DOCX = {RAW_DOCX!r};
      var gview = "https://docs.google.com/gview?url=" + encodeURIComponent(DOCX) + "&embedded=false";
      document.getElementById("btn-google").href = gview;

      var share = document.getElementById("share-url");
      if (location.hostname.indexOf("github.io") >= 0 || location.hostname.indexOf("jsdelivr") >= 0) {{
        share.value = location.href.split("?")[0];
      }}

      document.getElementById("copy").onclick = function () {{
        share.select();
        navigator.clipboard.writeText(share.value).then(function () {{
          document.getElementById("copied").classList.add("show");
        }});
      }};

      var q = new URLSearchParams(location.search);
      if (q.get("dl") === "1" || q.get("download") === "word") {{
        document.getElementById("btn-word").click();
      }}
      if (q.get("google") === "1") {{
        location.replace(gview);
      }}
    }})();
  </script>
</body>
</html>
"""


def set_run_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color=LINE, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_width(table, width_cm=17):
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def cell_para(cell, text, size=11, bold=False, color=INK, space_after=4):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_p(doc, text, size=11, bold=False, color=INK, space_after=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_h(doc, text):
    p = add_p(doc, text, size=16, bold=True, color=RED, space_after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RED)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    return add_p(doc, text, size=13, bold=True, color=RED2, space_after=6)


def grid(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, RED)
        set_cell_borders(cell, RED)
        set_cell_margins(cell)
        cell_para(cell, h, size=9, bold=True, color=WHITE)
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell_para(cell, val, size=10, bold=(c == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, SOFT)
    set_cell_borders(cell, RED)
    set_cell_margins(cell, 120, 120, 160, 160)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(label)
    set_run_font(r1, size=9, bold=True, color=RED)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2)

    cover = doc.add_table(rows=1, cols=1)
    set_table_width(cover)
    cell = cover.rows[0].cells[0]
    shade_cell(cell, DARK)
    set_cell_margins(cell, 220, 220, 240, 240)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("PROPOSTA COMERCIAL")
    set_run_font(r, size=9, bold=True, color="F5B5B8")
    p2 = cell.add_paragraph()
    r2 = p2.add_run("Empresa que Vende")
    set_run_font(r2, size=26, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    r3 = p3.add_run(f"{CLIENT}  ×  V4 Company")
    set_run_font(r3, size=14, color="FFD0D0")
    p4 = cell.add_paragraph()
    r4 = p4.add_run(
        "Consultoria, acompanhamento e execução assistida para transformar "
        "marketing e comercial em uma máquina de receita previsível — em 6 meses."
    )
    set_run_font(r4, size=11, color="E8D0D0")

    add_p(doc, "", space_after=6)
    grid(
        doc,
        ["Campo", "Informação"],
        [
            ["Cliente", CLIENT],
            ["Produto", "Empresa que Vende (EQV) — programa semestral de consultoria + execução assistida"],
            ["Prazo", "6 meses  ·  5 fases  ·  squad de 4 especialistas (Conteúdo, Mídia Paga, Comercial, CRM)"],
            ["Investimento", f"{PARCELAS}  ·  PIX à vista {AVISTA}  ·  referência {TABELA}"],
            ["Cashback EE", f"{CASHBACK} já investidos na Estruturação Estratégica, aplicados nesta condição"],
            ["Condução", "V4 Company  ·  consultor estratégico (1h/mês) + squad de execução assistida"],
            ["Incluso", "Comunidade, monitoria coletiva mensal, conteúdo semanal e Fábrica de Receita"],
            ["Garantia", "30 dias de serviço — reembolso incondicional se a entrega não fizer sentido"],
        ],
    )
    callout(
        doc,
        "LEITURA CENTRAL",
        "Você trabalha muito. A receita ainda não é previsível. O Empresa que Vende não executa "
        "no seu lugar: estrutura o negócio, treina o time e acompanha a execução. Quem opera é você.",
    )

    add_h(doc, "1. Objetivo")
    add_p(
        doc,
        f"Instalar, em 6 meses, o método para a {CLIENT} vender com previsibilidade: diagnóstico "
        "contínuo de maturidade GTM, consultoria estratégica mensal, execução assistida nas quatro "
        "frentes (conteúdo, mídia, comercial e CRM), comunidade de pares e imersão presencial da "
        "Fábrica de Receita.",
    )

    add_h(doc, "2. O problema que o programa resolve")
    add_p(
        doc,
        "O empresário típico do EQV já vende. O que falta é sistema para repetir o mês. Vendas "
        "oscilam; o processo mora na cabeça do melhor vendedor; marketing decide no achismo; a "
        "agência executa e o crescimento para quando o contrato termina.",
    )
    grid(
        doc,
        ["Sintoma", "Efeito na operação"],
        [
            ["Receita imprevisível", "O mês é esforço, não método. Sem processo, o resultado não se repete."],
            ["Processo na cabeça de uma pessoa", "Quando o melhor vendedor sai, o comercial sai junto."],
            ["Marketing no achismo", "Canal, conteúdo e mídia sem tese. Investe, testa, para."],
            ["Agência executa — você não aprende", "Quando o contrato acaba, o resultado vai junto."],
            ["Comercial amador", "Script, follow-up e pipeline improvisados. O funil vaza depois do lead."],
            ["Posicionamento frouxo", "A marca não escolheu ICP. O mercado escolhe preço."],
        ],
    )

    add_h(doc, "3. Jornada de 6 meses")
    grid(
        doc,
        ["Fase", "Quando", "Nome", "Função"],
        [
            ["01", "Mês 1", "Onboarding", "Kickoff, maturidade GTM e prioridades do ciclo."],
            ["02", "Mês 2", "EE 3.0 completa", "3P-IA, posicionamento ativo e canal de vendas no ar."],
            ["03", "Mês 3", "Consultoria de conteúdo", "Produção assistida + linha editorial no método."],
            ["04", "Mês 4", "Consultoria de tráfego", "Campanhas, pixel, audiência e otimização."],
            ["05", "Mês 5/6", "Consultoria comercial + CRM", "Scripts, follow-up, pipeline e monetização."],
        ],
    )

    add_h(doc, "4. Investimento")
    grid(
        doc,
        ["Linha", "Valor"],
        [
            ["Valor de tabela / referência", TABELA],
            ["PIX / à vista", AVISTA],
            ["Parcelado", f"{PARCELAS}  ·  total {TOTAL_PARCELADO}"],
            ["Cashback EE aplicado", CASHBACK],
            ["Economia no PIX vs. tabela", ECONOMIA],
        ],
    )
    callout(
        doc,
        "CONDIÇÃO COMERCIAL",
        f"PIX à vista {AVISTA}. Ou {PARCELAS} sem juros. Cashback de {CASHBACK} da "
        f"Estruturação Estratégica já aplicado sobre a referência de {TABELA}.",
    )
    add_h2(doc, "4.1 Garantia")
    add_p(
        doc,
        "Se após 30 dias de serviço você não estiver satisfeito com a entrega, devolvemos o "
        "investimento efetuado. Risco zero para entrar.",
    )

    add_h(doc, "5. Próximos passos")
    grid(
        doc,
        ["Responsável", "Ação"],
        [
            [CLIENT, f"Validar a condição comercial (PIX {AVISTA} ou {PARCELAS}) e formalizar o start."],
            ["V4", "Kickoff do mês 01: diagnóstico de maturidade GTM, prioridades e alocação do squad."],
            [f"V4 + {CLIENT}", "Entrar na comunidade, na monitoria mensal e no calendário da Fábrica de Receita."],
        ],
    )
    add_p(
        doc,
        f"Documento elaborado pela V4 Company para a {CLIENT}. Valores conforme condição comercial "
        f"da base ativa, com cashback de {CASHBACK} da EE aplicado sobre {TABELA}.",
        size=10,
        italic=True,
        color=MUT,
    )
    p = add_p(doc, f"{CLIENT} × V4 Company  ·  Empresa que Vende  ·  Confidencial", size=9, bold=True, color=RED)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print("Wrote", OUT_DOCX)


def build_html():
    html = html_doc()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_HTML.write_text(html, encoding="utf-8")
    OUT_DOC.write_text(html, encoding="utf-8")
    OUT_LANDING.write_text(landing_html(), encoding="utf-8")
    print("Wrote", OUT_HTML)
    print("Wrote", OUT_DOC)
    print("Wrote", OUT_LANDING)


if __name__ == "__main__":
    build_html()
    build_docx()
    print("Preview HTML:", PREVIEW)
    print("Word .doc:", RAW_DOC)
