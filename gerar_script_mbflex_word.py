#!/usr/bin/env python3
"""Script de pitch MBFlex em Word — fala didática, sem jargão."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RED = "E50914"
DARK = "280001"
INK = "1A1A1A"
MUT = "4A4A4A"
CREAM = "F7F5F4"
SOFT = "F8EFEF"
LINE = "E6D4D4"
WHITE = "FFFFFF"
GREEN = "1F6B3A"

OUT = Path("/workspace/script-mbflex-pitch.docx")


def set_run_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="E6D4D4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_width(table, width_cm):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True


def add_p(doc, text, size=11, bold=False, color=INK, space_before=0, space_after=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), RED)
        pBdr.append(bottom)
        pPr.append(pBdr)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=RED if level == 1 else DARK)
    return p


def add_bullet(doc, text, size=11, color=MUT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)
    return p


def box(doc, label, text, fill=SOFT, border=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 17.0)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, border, "12")
    set_cell_margins(cell, 100, 100, 140, 140)
    prevent_row_split(table.rows[0])
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(label.upper())
    set_run_font(r1, size=9, bold=True, color=border)
    for i, para in enumerate(text.split("\n")):
        p = p1 if i == 0 and not p1.runs else cell.add_paragraph()
        if i == 0 and p1.runs:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(para)
        set_run_font(r, size=11, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def fala_box(doc, text):
    """Spoken script — the part Sara reads almost as-is."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 17.0)
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF7F7")
    set_cell_borders(cell, RED, "16")
    set_cell_margins(cell, 110, 110, 160, 160)
    prevent_row_split(table.rows[0])
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run("O QUE FALAR")
    set_run_font(r1, size=9, bold=True, color=RED)
    for i, para in enumerate([p.strip() for p in text.strip().split("\n\n") if p.strip()]):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(para)
        set_run_font(r, size=12, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def slide_head(doc, num, title, tempo):
    add_heading(doc, f"{num}  ·  {title}", 1)
    add_p(doc, tempo, size=10, italic=True, color=MUT, space_after=8)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("MBFlex Express  ×  V4  ·  Script interno  ·  não enviar ao cliente")
    set_run_font(r, size=8, color=RED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = fp.add_run("Assessoria de Growth  ·  11 slides  ·  Página ")
    set_run_font(r, size=8, color=MUT)
    add_page_number(fp)

    # capa
    cover = doc.add_table(rows=1, cols=1)
    set_table_width(cover, 17.0)
    cell = cover.cell(0, 0)
    shade_cell(cell, DARK)
    set_cell_borders(cell, DARK, "4")
    set_cell_margins(cell, 280, 280, 220, 220)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("SCRIPT DE CONVERSA")
    set_run_font(r, size=11, bold=True, color="FFB4A0")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    r = p2.add_run("MBFlex Express")
    set_run_font(r, size=28, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    r = p3.add_run("Como explicar a proposta — com clareza, sem enrolar")
    set_run_font(r, size=14, color="FFEBC8")
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(10)
    r = p4.add_run("Márcio e Vinícius  ·  11 slides  ·  projeto 12 meses  ·  12× R$ 9.259,31")
    set_run_font(r, size=11, color="FFB4A0")
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    box(
        doc,
        "Como usar este Word",
        "A caixa vermelha é o que você fala. Está escrita para um cliente prático, que não compra jargão.\n"
        "As caixas cinza são só para você: o que veio da reunião, o que não falar, a conta do preço.\n"
        "Uma ideia por vez. Se ele franzir a testa, para e explica com o exemplo do próprio negócio — Mercado Livre, Shopee, motoboy, primeira coleta.\n"
        "Não leia o slide. Conversa. Preço só no final.",
    )

    add_heading(doc, "Antes de ligar a tela", 1)
    add_p(
        doc,
        "Márcio é sócio e comercial. Gosta de conta que fecha. Não gosta de fantasia. Vinícius é financeiro — pensa caixa, qualidade e se a operação aguenta. Rafael é operação: se entrar na call, não puxa ele para decidir o projeto.",
        space_after=8,
    )
    add_p(doc, "Eles já pagaram a Estruturação (R$ 14.612). Esta call não vende isso de novo. Vende o time que opera o canal, todo mês, por 12 meses.", space_after=8)

    fala_box(
        doc,
        "Antes de eu abrir a tela, quero ouvir vocês.\n\n"
        "Como está saindo essa fase da Estruturação? O que mais ficou claro para vocês até agora?\n\n"
        "Pode ser a dependência das plataformas, o teto da operação, ou o fato de ainda não terem um caminho próprio para trazer loja — o que vocês quiserem colocar na mesa.",
    )
    box(
        doc,
        "Só você · o que ele já disse na reunião (não cita como slide)",
        "“Preciso aquecer o meu comercial, preciso ser uma máquina de captar cliente.”\n"
        "“Hoje tá vindo tudo orgânico. A gente não vai atrás de nenhum cliente.”\n"
        "“Eu quero qualidade e não quantidade. Com quantidade no nosso setor a gente se bagaça todo.”\n"
        "Se ele repetir a frase da máquina: “É isso — e hoje essa máquina é a plataforma.” Sem jogar a frase dele de volta como citação.",
        fill=CREAM,
        border=DARK,
    )

    add_heading(doc, "Palavras simples (se ele perguntar “o que é isso?”)", 1)
    add_p(doc, "Não ensina dicionário. Se precisar, usa estes equivalentes — uma frase, e segue.", space_after=6)
    rows = [
        ("Marketplace / plataforma", "Mercado Livre, Shopee, Amazon. Eles indicam o cliente. Vocês não escolhem quem chega."),
        ("Loja própria", "O lojista que vende no site dele, não só dentro da plataforma. Esse cliente a plataforma não manda."),
        ("Malha", "A operação de vocês: motoboy, região, capacidade de entregar no mesmo dia."),
        ("Região com capacidade", "Bairro/zona onde vocês já têm entregador e conseguem cumprir o prazo. Sem isso, anúncio só enche o WhatsApp de pedido que não cabe."),
        ("Primeira coleta", "O cliente de verdade: vocês foram lá, pegaram o pacote. Não é clique, não é conversa. É receita."),
        ("Mix 50/50", "Metade da carteira vinda de loja própria, metade continuando nas plataformas. Diversificar — não abandonar o ML."),
        ("Assessoria de Growth", "Um time da V4 na conta, todo mês: quem conversa com vocês, quem cuida do Google, quem escreve e quem desenha."),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    set_table_width(table, 17.0)
    for i, h in enumerate(["Na conversa", "Como explicar em uma frase"]):
        c = table.rows[0].cells[i]
        shade_cell(c, RED)
        set_cell_borders(c, RED, "4")
        set_cell_margins(c, 70, 70, 90, 90)
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color=WHITE)
    for ri, (a, b) in enumerate(rows):
        for ci, val in enumerate((a, b)):
            c = table.rows[ri + 1].cells[ci]
            shade_cell(c, CREAM if ri % 2 else WHITE)
            set_cell_borders(c, LINE, "4")
            set_cell_margins(c, 70, 70, 90, 90)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=10, bold=ci == 0, color=INK if ci == 0 else MUT)
        prevent_row_split(table.rows[ri + 1])
    prevent_row_split(table.rows[0])
    table.rows[0].cells[0].width = Cm(4.6)
    table.rows[0].cells[1].width = Cm(12.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    box(
        doc,
        "Nunca na boca — mesmo que esteja na sua cabeça",
        "CRM  ·  “vocês nunca investiram em marketing”  ·  Paulo / Renata  ·  “desconto”  ·  “fee”  ·  cartão  ·  “60/40 do P&L”  ·  “a gente chega em 500 com anúncio”  ·  vender Estruturação de novo  ·  misturar lojista e motoboy no mesmo anúncio.",
        fill="FFF4E5",
        border="C2410C",
    )

    # 01
    slide_head(doc, "SLIDE 01", "Capa — da plataforma para um caminho próprio", "20 a 30 segundos  ·  institucional")
    add_p(doc, "O slide já diz a tese. Você só traduz. Sem abrir número. Sem “máquina de captação”.", size=11, color=MUT, space_after=6)
    fala_box(
        doc,
        "A gente montou esta conversa em cima do que já foi estudado com vocês — não em cima do site.\n\n"
        "A frase que resume tudo: a operação de vocês já existe. O caminho para trazer loja por conta própria, não.\n\n"
        "Sem esse caminho, quem manda na carteira continua sendo o Mercado Livre, a Shopee e a Amazon.\n\n"
        "Eu vou primeiro pelos números e por essa dependência. Solução e valor vêm no fim. Sem pressa.",
    )
    box(doc, "Pergunta", "“Está ok a gente começar pelos números da operação, antes de falar de time e de valor?”", fill="F0FDF4", border=GREEN)

    # 02
    slide_head(doc, "SLIDE 02", "Dependência — não é problema de operação", "cerca de 2 minutos")
    add_p(doc, "Três cartões no slide: de onde vem o cliente, loja própria zerada, e o teto da malha. Fecha na leitura de baixo. Não lê parágrafo por parágrafo.", size=11, color=MUT)
    fala_box(
        doc,
        "Vocês não têm problema de operação. A malha entrega. Cinco estrelas. Cliente fica.\n\n"
        "O que vocês têm é uma carteira amarrada nas plataformas. Mercado Livre, Shopee e Amazon indicam. O cliente chega. Mas a origem não é de vocês — é da plataforma.\n\n"
        "Loja que vende no site dela, fora do marketplace, hoje é zero na carteira. Esse público a plataforma não traz. Sem um caminho próprio, continua 100% na mão deles.\n\n"
        "E a operação já mostrou o limite: no fim do ano recusaram mais de 200 clientes, porque estava lotado. E perderam um entregador forte — foi para quem tinha mais pacote, mais constância.\n\n"
        "Em uma frase: a operação entrega. O comercial não escolhe quem entra. Quem escolhe é preço e regra da plataforma.\n\n"
        "O motoboy é a maior conta de vocês. Pedido errado esvazia a rua. Pedido certo segura o entregador.\n\n"
        "Chegar em 500 do mesmo jeito — só com indicação de plataforma — aprofunda essa dependência. A decisão não é acelerar a indicação. É diversificar de onde o cliente vem.",
    )
    box(
        doc,
        "Se ele perguntar “o que é loja própria?”",
        "“É o lojista que tem o site dele, a venda dele, e precisa de same day. Não é o seller que só vive dentro do Mercado Livre. A plataforma não indica esse cara. Por isso precisa de um caminho de vocês.”",
        fill=CREAM,
        border=DARK,
    )
    box(doc, "Pergunta", "“Se os próximos clientes vierem do mesmo jeito que os 196 de hoje — só pela plataforma — o que mais preocupa vocês: a dependência, ou o teto da operação?”", fill="F0FDF4", border=GREEN)

    # 03
    slide_head(doc, "SLIDE 03", "Os números — 196 de um lado, zero do outro", "cerca de 2 minutos")
    fala_box(
        doc,
        "Do lado que já existe: cerca de 196 clientes, em torno de um milhão por mês, sete mil pacotes por dia, uns 140 entregadores, mais de 100 mil entregas no mês. Oitenta por cento moto, vinte por cento carro. Cinco estrelas no Google. Cliente não foge.\n\n"
        "Isso é operação de verdade. Não é empresa começando.\n\n"
        "Do lado que não existe: zero por cento da carteira vem de loja própria. Cem por cento chega por indicação de marketplace. Ainda não tem um caminho para o lojista que busca vocês no Google. Anúncio ainda não trabalha isso.\n\n"
        "A meta de vocês no ano é 500. De 196 para 500 faltam 304.\n\n"
        "A pergunta honesta é: de onde vêm esses 304? Se a resposta for “a plataforma indica”, o desenho não muda. Continuam 100% na mão do marketplace — só que com mais gente. E 500 vira mais do mesmo.",
    )
    box(
        doc,
        "Só você · não cravar ticket de R$ 10 mil",
        "Na estratégia interna aparece cerca de R$ 10 mil por cliente. 10 mil × 196 daria quase R$ 2 milhões — e o faturamento que eles informaram é cerca de R$ 1 milhão. Márcio pega conta torta.\n"
        "Se o Vinícius puxar: “Esse valor por cliente a gente crava no início do trabalho, com calma. Hoje eu não vou usar um número que ainda precisa bater com o milhão do mês.”",
        fill=CREAM,
        border=DARK,
    )
    box(doc, "Pergunta", "“Se eu te perguntar de onde vêm os próximos 304 — a plataforma, ou um caminho de vocês — o que vocês responderiam hoje?”", fill="F0FDF4", border=GREEN)

    # 04
    slide_head(doc, "SLIDE 04", "O que precisa existir antes de anunciar", "cerca de 1 minuto e meio")
    add_p(doc, "Seis pontos no slide. Passe como lista falada. O fechamento é: anunciar em cima disso bagunça a malha.", size=11, color=MUT)
    fala_box(
        doc,
        "Antes de gastar real em anúncio, tem seis coisas que precisam estar claras. Eu vou passar uma a uma, no português da operação.\n\n"
        "Um. Hoje quem chega é quem a plataforma manda. Vocês não escolhem o cliente.\n\n"
        "Dois. Não existe ainda um caminho para o lojista de site próprio. Não tem busca no Google, não tem página falando com ele, não tem rotina.\n\n"
        "Três. O que cabe ou não cabe — geladeira, vaso, região — ainda se decide na conversa do WhatsApp. Se anúncio entrar sem esse filtro, entra volume errado.\n\n"
        "Quatro. Sete mil pacotes por dia é o teto. Sem região com entregador sobrando, Google não sobe. Qualidade cinco estrelas vem primeiro.\n\n"
        "Cinco. São dois públicos: o lojista e o motoboy. Se os dois caem no mesmo anúncio ou no mesmo WhatsApp, estraga conversão e estraga reputação.\n\n"
        "Seis. A meta de 500, sem um caminho de loja própria, é só mais volume na mesma dependência.\n\n"
        "Por isso a gente não liga anúncio amanhã. Ligar anúncio em cima disso é pagar para bagunçar a malha.",
    )
    box(
        doc,
        "Se ele disser “então vamos anunciar já”",
        "“Eu entendo a pressa. O ponto é: anúncio sem filtro e sem região vira pedido que vocês vão recusar de novo — ou motoboy que não segura. A gente liga o Google quando a região aguenta. É exatamente para não repetir o fim de ano.”",
        fill=CREAM,
        border=DARK,
    )
    box(doc, "Pergunta", "“Dessas seis, se a gente ligasse anúncio amanhã, qual vocês acham que quebraria primeiro?”", fill="F0FDF4", border=GREEN)

    # 05
    slide_head(doc, "SLIDE 05", "A promessa — duas frases", "20 segundos  ·  não alongar")
    fala_box(
        doc,
        "O que a V4 faz daqui para frente: opera o caminho de loja própria que a plataforma não traz.\n\n"
        "Não é um PDF. É o time que liga esse caminho, todo mês.",
    )
    box(doc, "Só você", "Não explique produto aqui. Time vem no próximo. Não fale “a Estruturação que vocês já pagaram é o mapa”.", fill=CREAM, border=DARK)

    # 06
    slide_head(doc, "SLIDE 06", "O time — quem entra na conta", "cerca de 2 minutos  ·  não pular")
    add_p(doc, "Quatro pessoas. Fala o que cada uma faz na vida real da MBFlex, não o cargo.", size=11, color=MUT)
    fala_box(
        doc,
        "Quem entra na conta, todo mês, são quatro pessoas. Não é um fornecedor de anúncio. É um time.\n\n"
        "A Account é quem conversa com você e com o Vinícius. Ela cuida do equilíbrio: metade loja própria, metade plataforma. Ela também segura o freio: anúncio só sobe se a região aguentar. E ela acompanha até a primeira coleta — não até o clique.\n\n"
        "O gestor de tráfego cuida do Google: quem já está procurando entrega no mesmo dia, na região de vocês. Instagram e Facebook, neste desenho, só para quem já conheceu e não fechou. Sem região liberada, a verba não sobe.\n\n"
        "Quem escreve o texto fala a língua da operação: loja própria, cinco estrelas, o que cabe na moto. Motoboy tem conversa separada — não entra no mesmo anúncio do lojista.\n\n"
        "Quem desenha deixa anúncio, página e botão com a mesma cara. Sem aquele “Flex SP” genérico que parece todo mundo.\n\n"
        "A Estruturação já deixou o mapa claro. Growth é quem opera isso todo mês.",
    )
    box(
        doc,
        "Se o Vinícius puxar motoboy",
        "“A gente olha o entregador, sim — em um caminho separado. Mesmo cuidado, conversa diferente. Lojista e motoboy no mesmo anúncio estraga os dois. Volume certo na região é o que segura o motoboy na rua.”",
        fill=CREAM,
        border=DARK,
    )
    box(
        doc,
        "Pergunta (mata “quem faz?” antes do preço)",
        "“Olhando essas quatro pessoas no dia a dia de vocês — o que ficou mais claro? Tem alguma que vocês querem que eu explique melhor antes de a gente seguir?”",
        fill="F0FDF4",
        border=GREEN,
    )

    # 07
    slide_head(doc, "SLIDE 07", "Como isso roda — Google, depois conversa, depois coleta", "cerca de 2 minutos")
    fala_box(
        doc,
        "Vou desenhar o caminho com as mãos, porque é simples.\n\n"
        "O lojista pesquisa no Google alguma coisa do tipo “entrega no mesmo dia na minha região”. Cai numa página de vocês, que já filtra o que cabe na moto. Se fizer sentido, vocês conversam. Se fechar, vocês fazem a primeira coleta. Aí a gente sabe que aquele anúncio trouxe cliente de verdade — não só conversa no WhatsApp.\n\n"
        "Isso é um projeto de doze meses. Os três primeiros meses são só para ligar o caminho com segurança.\n\n"
        "Mês 1: organizar. Quais regiões aguentam. O que a página vai dizer. O que não entra — geladeira, linha branca, o que vocês já recusam. Sem isso, anúncio não sobe.\n\n"
        "Mês 2: um teste pequeno no Google. A gente descobre quanto custa trazer um cliente até a primeira coleta — não até o clique.\n\n"
        "Mês 3: quem visitou e não fechou, a gente volta a falar. Instagram e Facebook entram aqui, para recuperar — não para caçar gente fria.\n\n"
        "Do mês 4 ao 12 o time continua. Não é um projeto de três meses. Três meses é só o começo com o freio puxado.",
    )
    box(
        doc,
        "Se perguntar “por que 90 dias?”",
        "“Não é um prazo de vocês e não é garantia de 500 clientes. É o tempo para ligar o caminho sem bagunçar a malha. O contrato é de 12 meses. Depois do terceiro mês a gente continua operando — com mais leitura, não começando do zero.”",
        fill=CREAM,
        border=DARK,
    )
    box(doc, "Pergunta", "“Esse ritmo — primeiro organizar, depois um teste no Google, e só então anúncio de recuperação — é o ritmo que vocês querem para não repetir o fim de ano?”", fill="F0FDF4", border=GREEN)

    # 08
    slide_head(doc, "SLIDE 08", "O que a gente olha junto", "cerca de 1 minuto")
    fala_box(
        doc,
        "O que a gente vai olhar com vocês, mês a mês, é bem concreto.\n\n"
        "Se o anúncio virou primeira coleta — ou só conversa.\n"
        "Quanto da carteira já é loja própria, e quanto ainda é plataforma.\n"
        "Quanto custa um cliente que de fato começou a usar vocês — não quanto custa um clique.\n"
        "Google só no ar onde a região aguenta.\n"
        "Motoboy em conversa separada da do lojista.\n"
        "A Account no ritmo com você e com o Vinícius.\n\n"
        "A meta de 500 é da empresa. O papel deste trabalho é ligar a metade que hoje é zero: loja própria. A plataforma continua indicando. A gente não vai tratar anúncio como se ele sozinho trouxesse os 304.",
    )
    box(
        doc,
        "Só você · a planilha não prova 500",
        "Com R$ 3.500 de mídia, a conta interna chega perto de 27 lojas no ano — não 304. Se ele perguntar se isso entrega 500: “Não. 500 é a meta de vocês no ano, somando plataforma e loja própria. Este trabalho liga o pedaço que hoje não existe.”",
        fill=CREAM,
        border=DARK,
    )

    # 09
    slide_head(doc, "SLIDE 09", "A operação de vocês no desenho", "cerca de 1 minuto e meio")
    fala_box(
        doc,
        "Isso não é um pacote genérico de agência. É a operação de vocês desenhada.\n\n"
        "Dois públicos, duas conversas: o lojista e o motoboy. Mesmo cuidado, caminhos separados. Os dois no mesmo anúncio ou no mesmo WhatsApp não funciona.\n\n"
        "A gente não vai reabrir o diagnóstico. O que precisa ser feito já está claro. Agora é ligar o caminho de loja própria em cima do que vocês já definiram.\n\n"
        "O que este trabalho cria: um jeito de o lojista de site próprio achar vocês, conversar, e virar primeira coleta.\n\n"
        "O que a plataforma continua fazendo: indicar seller, manter a base de vocês, manter as cinco estrelas. Os outros 50% da carteira seguem aí. A gente não mexe nisso para pior.",
    )
    box(doc, "Pergunta", "“Esse recorte está do tamanho certo? A gente cria o que falta. A plataforma continua com o que já funciona.”", fill="F0FDF4", border=GREEN)
    box(doc, "Não falar", "“O Vinícius pediu.”  ·  “A Estruturação já paga o mapa.”  ·  “Não é um pacote genérico —” (o slide já limpou isso).", fill="FFF4E5", border="C2410C")

    # 10
    slide_head(doc, "SLIDE 10", "Recap — antes do valor", "pausa de 5 segundos, depois uma frase")
    fala_box(
        doc,
        "Em uma frase, para a gente olhar no mesmo quadro:\n\n"
        "A operação existe. A carteira depende da plataforma. Loja própria é zero. O time liga esse caminho. A meta de 500 é de vocês — não é uma promessa de anúncio.\n\n"
        "Se isso estiver alinhado, o próximo passo é a condição do projeto de 12 meses.",
    )
    box(
        doc,
        "Não vira o slide 11 ainda",
        "Se Márcio quiser falar de motoboy, mix ou região — fica. Preço só quando o desenho estiver quieto.",
        fill=CREAM,
        border=DARK,
    )

    add_heading(doc, "Antes do preço — quatro perguntas", 1)
    add_p(doc, "Quando o valor entrar, a única discussão que pode restar é a condição. Aqui você mata o resto.", space_after=6)
    fala_box(
        doc,
        "Antes de eu trazer o valor, quero alinhar quatro coisas — para a gente não chegar no número com dúvida aberta.\n\n"
        "O que vocês mais gostaram da leitura até aqui?\n\n"
        "O que vocês acham que a gente não pode deixar de fazer?\n\n"
        "Quando faz sentido começar o mês de organização — o primeiro mês, ainda sem anúncio grande?\n\n"
        "Vocês querem a V4 operando esse próximo passo, ou estão olhando outro caminho?",
    )
    box(
        doc,
        "Se Márcio pedir o número cedo",
        "“Eu te trago a conta no fim, porque ela só fecha se o desenho estiver claro. Me dá mais dois minutos.” Se insistir, vai ao 11 — e não reabre escopo em cima do valor.",
        fill=CREAM,
        border=DARK,
    )

    # 11
    slide_head(doc, "SLIDE 11", "Investimento — projeto de 12 meses", "cerca de 2 minutos  ·  último slide")
    add_p(doc, "Fala devagar. Âncora primeiro. Depois o que já investiram. Depois o número grande. Respira. Ele vai conferir a conta.", size=11, color=MUT)
    fala_box(
        doc,
        "O que estamos contratando: um projeto de 12 meses. Assessoria de Growth. As quatro pessoas que a gente viu, todo mês.\n\n"
        "O valor de tabela desse time é R$ 11.000 por mês.\n\n"
        "Vocês já investiram R$ 14.612 na Estruturação. A gente reconhece isso na mensalidade — dá R$ 1.217,67 por mês — e entra uma condição comercial de R$ 523,02.\n\n"
        "Fica 12 vezes R$ 9.259,31 ao mês. É valor mensal, projeto de um ano.\n\n"
        "Dinheiro de anúncio é à parte. A conta de partida é R$ 3.500 por mês. Esse valor não sobe se a região não aguentar e se a gente não estiver vendo primeira coleta de verdade.",
    )

    math = doc.add_table(rows=5, cols=2)
    set_table_width(math, 17.0)
    lines = [
        ("Valor de tabela", "R$ 11.000,00 / mês"),
        ("Estruturação reconhecida (14.612 ÷ 12)", "− R$ 1.217,67"),
        ("Condição comercial", "− R$ 523,02"),
        ("Valor mensal do projeto", "R$ 9.259,31"),
        ("Projeto 12 meses", "12 × R$ 9.259,31"),
    ]
    for i, (a, b) in enumerate(lines):
        for ci, val in enumerate((a, b)):
            c = math.rows[i].cells[ci]
            fill = RED if i >= 3 else (CREAM if i % 2 else WHITE)
            shade_cell(c, fill)
            set_cell_borders(c, RED if i >= 3 else LINE, "4")
            set_cell_margins(c, 80, 80, 100, 100)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=11, bold=True, color=WHITE if i >= 3 else INK)
        prevent_row_split(math.rows[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    box(
        doc,
        "Nunca na boca neste slide",
        "“Desconto.”  ·  “Fee.”  ·  “Cartão.”  ·  “Alinhada com o Vinícius.”  ·  “Isso se paga em X meses” com número da planilha.\n"
        "Fala: valor de tabela · reconhece o que já investiram · condição comercial · projeto de 12 meses.",
        fill="FFF4E5",
        border="C2410C",
    )
    box(
        doc,
        "Pergunta de avanço (não “faz sentido?”)",
        "“Olhando essa condição, o que vocês precisam alinhar para a gente avançar no projeto de 12 meses?”\n"
        "Ou: “Se o desenho fechou, a gente consegue decidir o projeto nesta conversa e deixar o valor de anúncio combinado nesta semana?”",
        fill="F0FDF4",
        border=GREEN,
    )

    add_heading(doc, "Pedir a venda", 1)
    fala_box(
        doc,
        "Eu quero conduzir esse próximo passo com vocês dois.\n\n"
        "A gente começa organizando região e página, testa o Google com o freio puxado, e segue 12 meses com a Account no ritmo, olhando primeira coleta — não clique.\n\n"
        "Se o plano fechou: avançamos o contrato e já marcamos o início.",
    )
    box(
        doc,
        "Só você · o close dele na reunião",
        "Ele disse: “Eu preciso que você me ajude a dar certo.” Não cita. Responde no tom: conta que fecha, malha no freio, time na operação.",
        fill=CREAM,
        border=DARK,
    )
    box(
        doc,
        "Pedir",
        "“Podemos avançar com o projeto de 12 meses nessa condição e marcar o início?”\n"
        "Se hesitar: “O que ainda está em aberto — o valor, a data de começar, ou alguma peça do desenho?”",
        fill="F0FDF4",
        border=GREEN,
    )

    add_heading(doc, "Se ele travar — respostas curtas", 1)
    objs = [
        ("“Isso entrega os 500?”", "Não sozinho. 500 é a meta de vocês no ano. Este trabalho liga a parte que hoje é zero: loja própria. A plataforma continua indicando."),
        ("“Por que 90 dias / três meses?”", "É só o começo com segurança, dentro de um projeto de 12 meses. Não é prazo de 500 clientes."),
        ("“Está caro.” / ele pega a calculadora", "Tabela 11 mil. Menos o que já investiram na Estruturação, mês a mês. Menos a condição. Fecha 9.259,31. Pode conferir: 11.000 menos 1.217,67 menos 523,02."),
        ("“Preciso pensar.”", "Claro. Pensar em quê — valor, data de começar, ou alguma peça do desenho? Assim eu te ajudo no ponto certo."),
        ("“E o motoboy?”", "Caminho separado. Volume certo na região é o que segura ele. Não mistura com o anúncio do lojista."),
        ("“Vamos anunciar já.”", "Anúncio sem região e sem filtro vira pedido que vocês recusam de novo. A gente liga o Google quando a malha aguenta."),
        ("“Como a gente acompanha?”", "A Account no ritmo com você e com o Vinícius. Mix na mesa. Custo até a primeira coleta. Anúncio só onde a região libera."),
        ("Ticket / “quanto o cliente paga?”", "Esse número a gente crava no início do trabalho, batendo com o milhão do mês. Hoje eu não uso um valor que ainda não fechou."),
    ]
    ot = doc.add_table(rows=1 + len(objs), cols=2)
    set_table_width(ot, 17.0)
    for i, h in enumerate(["Se ele disser", "Você responde"]):
        c = ot.rows[0].cells[i]
        shade_cell(c, DARK)
        set_cell_borders(c, DARK, "4")
        set_cell_margins(c, 70, 70, 90, 90)
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color=WHITE)
    for ri, (a, b) in enumerate(objs):
        for ci, val in enumerate((a, b)):
            c = ot.rows[ri + 1].cells[ci]
            shade_cell(c, CREAM if ri % 2 else WHITE)
            set_cell_borders(c, LINE, "4")
            set_cell_margins(c, 80, 80, 90, 90)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=10, bold=ci == 0, color=INK)
        prevent_row_split(ot.rows[ri + 1])
    prevent_row_split(ot.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading(doc, "Em uma frase, para não perder o fio", 1)
    box(
        doc,
        "Leva isto na cabeça",
        "Estudo de conta no começo. Uma ideia por vez. Anúncio só quando a região aguenta. 500 é meta deles, não promessa sua. Um projeto de 12 meses. 12 vezes R$ 9.259,31. Pede a venda com calma.",
        fill="F0FDF4",
        border=GREEN,
    )
    add_p(
        doc,
        "Fontes desta conversa: reunião de 22/07 com Márcio e Vinícius · estudo da conta · estratégia de 18/08 · Estruturação R$ 14.612 · planilha interna (não abrir na tela).",
        size=9,
        color=MUT,
        italic=True,
        space_before=8,
    )

    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
