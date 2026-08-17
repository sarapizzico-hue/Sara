#!/usr/bin/env python3
"""Gera o documento Word consolidado da Estruturação Comercial G6 Internet."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu, Twips

RED = "A10F14"
RED2 = "E50914"
INK = "1A1A1A"
MUT = "4A4A4A"
CREAM = "F7F5F4"
LINE = "E6D4D4"
WHITE = "FFFFFF"
SOFT = "F8EFEF"
OK = "167846"
DARK = "280001"

OUT = "/workspace/documento-g6-internet-estruturacao-comercial.docx"


def set_run_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="E6D4D4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_width(table, width_cm):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def cell_para(cell, text, size=11, bold=False, color=INK, align="left", space_after=4, name="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold, color=color)
    return p


def add_cell_run(cell, text, size=11, bold=False, color=INK, italic=False):
    p = cell.paragraphs[0] if cell.paragraphs[0].text == "" and len(cell.paragraphs) == 1 else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_p(doc, text, size=11, bold=False, color=INK, space_before=0, space_after=8, align="left", italic=False, name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), RED)
        pBdr.append(bottom)
        pPr.append(pBdr)
    run = p.add_run(text)
    set_run_font(run, name="Calibri", size=18 if level == 1 else 14, bold=True, color=RED if level == 1 else RED2)
    return p


def add_bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True, color=INK)
        r2 = p.add_run(text)
        set_run_font(r2, size=size, color=MUT)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=MUT)
    return p


def callout(doc, label, text, fill=SOFT, label_color=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 16.5)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, RED, "12")
    set_cell_margins(cell, 100, 100, 140, 140)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(label.upper())
    set_run_font(r1, size=9, bold=True, color=label_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, bold=False, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def make_table(doc, headers, rows, col_widths=None, header_fill=RED):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 16.5)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_borders(cell, header_fill, "4")
        set_cell_margins(cell, 70, 70, 90, 90)
        cell_para(cell, h, size=9, bold=True, color=WHITE, align="left", space_after=0)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            fill = CREAM if r_i % 2 else WHITE
            shade_cell(cell, fill)
            set_cell_borders(cell, LINE, "4")
            set_cell_margins(cell, 70, 70, 90, 90)
            cell_para(cell, val, size=10, bold=c_i == 0, color=INK, space_after=0)
        prevent_row_split(table.rows[r_i + 1])
    prevent_row_split(table.rows[0])
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)


def build():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.0)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("G6 Internet  ×  V4 Company  ·  Estruturação Comercial")
    set_run_font(r, size=8, color=RED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = fp.add_run("Documento consolidado  ·  Confidencial  ·  Página ")
    set_run_font(r, size=8, color=MUT)
    add_page_number(fp)
    r2 = fp.add_run("  ·  Reunião 14/08/2026")
    set_run_font(r2, size=8, color=MUT)

    # ========== CAPA ==========
    cover = doc.add_table(rows=1, cols=1)
    set_table_width(cover, 17.0)
    c = cover.cell(0, 0)
    shade_cell(c, DARK)
    set_cell_borders(c, DARK, "0")
    set_cell_margins(c, 280, 280, 220, 220)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("G6 INTERNET  ×  V4 COMPANY")
    set_run_font(r, size=11, bold=True, color="F5B5B8")
    p2 = c.add_paragraph()
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run("Estruturação Comercial")
    set_run_font(r, name="Calibri", size=32, bold=True, color=WHITE)
    p3 = c.add_paragraph()
    p3.paragraph_format.space_after = Pt(14)
    r = p3.add_run("Do comercial artesanal à máquina de receita")
    set_run_font(r, size=16, color="FFD0D0", italic=True)
    p4 = c.add_paragraph()
    r = p4.add_run(
        "Documento consolidado da proposta apresentada e da reunião comercial "
        "de 14 de agosto de 2026 — diagnóstico, escopo, fases, investimento, "
        "decisões e próximos passos."
    )
    set_run_font(r, size=12, color="E8D0D0")

    add_p(doc, "", size=8, space_after=6)

    meta = [
        ["Cliente", "G6 Internet"],
        ["Produto", "Estruturação Comercial (Sales V4) — projeto fechado com implementação assistida"],
        ["Duração", "12 semanas  ·  106,5 horas  ·  dual-track loja × PAP"],
        ["Investimento", "R$ 71.582,32  (referência do produto: R$ 89.635,85)  ·  12× recorrente"],
        ["Reunião", "14/08/2026  ·  [Comercial] V4 Company :::: G6 Internet"],
        ["Participantes", "Laurelise Santos (G6)  ·  Rafaela Tolomeu Messias  ·  Sara Pizzico  ·  Denis Orosco"],
        ["Pendência", "Validação da proposta com o sócio Geraldo antes do fechamento"],
        ["Fontes", "Deck de 18 slides da Estruturação Comercial + ata/transcrição da reunião (Google Docs)"],
    ]
    make_table(doc, ["Campo", "Informação"], meta, col_widths=[4.2, 12.3])

    callout(
        doc,
        "Tese da apresentação",
        "A G6 não tem problema de lead. Tem um comercial reativo sem sistema. "
        "Entra cerca de 1.000 clientes por mês e a base quase não cresce. "
        "A Estruturação Comercial instala a fundação — processo, CRM em uso, gestão e rotina — "
        "para o novo gestor (Mateus) herdar uma máquina pronta, e não improvisar processo, peça e meta ao mesmo tempo.",
    )

    # ========== SUMÁRIO ==========
    add_heading_styled(doc, "Sumário", 1)
    toc = [
        "1.  Resumo executivo",
        "2.  Contexto da G6 — o que foi diagnosticado",
        "3.  A conta do vazamento — dados da operação e da reunião",
        "4.  Dores × o que a Estruturação Comercial fecha",
        "5.  Visão estratégica e promessa do projeto",
        "6.  Formato, escopo e o que fica instalado",
        "7.  As seis fases (com entregas e encaixe G6)",
        "8.  Três pilares: sistemas, processos e pessoas",
        "9.  Critérios de Receita Previsível",
        "10. Investimento, condição comercial e viabilidade",
        "11. Ata da reunião de 14/08/2026 — decisões e próximos passos",
        "12. Pauta complementar: lançamento G6 Móvel / Chip",
        "13. Como levar este caso ao Geraldo",
        "14. Fontes e materiais de apoio",
    ]
    for item in toc:
        add_p(doc, item, size=12, color=INK, space_after=4)

    # ========== 1 ==========
    add_heading_styled(doc, "1. Resumo executivo", 1)
    add_p(
        doc,
        "Em 14 de agosto de 2026, a V4 Company apresentou à G6 Internet a proposta de "
        "Estruturação Comercial: um projeto de 12 semanas, com 106,5 horas, para construir "
        "e instalar a fundação comercial da operação — loja e porta a porta (PAP) — com "
        "diagnóstico presencial, arquitetura, CRM como painel de receita, camada estratégica, "
        "imersão/treino e handover de 60 dias.",
        align="justify",
    )
    add_p(
        doc,
        "A leitura central: a G6 já gera demanda (cerca de 1.000 ativações/mês, WhatsApp em ~1 minuto, "
        "Meta Forms com 1.269 leads no ano e CPL de R$ 7,51). O que falta é sistema. O CRM existe, "
        "mas não governa. Loja e PAP vendem sem padrão. A meta desce quando aperta. Ativação, save "
        "e retenção ficam fora do processo. Entre janeiro e junho de 2026, a operação somou 3.755 cancelamentos.",
        align="justify",
    )
    add_p(
        doc,
        "Na reunião, Rafaela traduziu esse vazamento em conta de gestão: cada cliente custa cerca de "
        "R$ 800 (CAC + instalação) e precisa ficar no mínimo 8 meses para se pagar. Sobre os 3.755 "
        "cancelamentos, a conta de padaria aponta ~R$ 3 milhões já investidos para adquirir/instalar "
        "e ~R$ 4 milhões de receita potencial deixada na mesa — cerca de R$ 7 milhões em seis meses. "
        "O investimento de marketing com a V4 + a Estruturação Comercial não chega a 10% desse vazamento.",
        align="justify",
    )
    add_p(
        doc,
        "Laurelise (Laura) reconheceu o encaixe com a dor combinada na conversa anterior, afirmou "
        "vontade de fazer o projeto e deixou pendente a validação com o sócio Geraldo. O cronograma "
        "de 12 semanas e o diagnóstico in loco (18h) foram alinhados. O valor da condição comercial "
        "G6 é R$ 71.582,32, parcelável em 12 recorrências (sem cartão), para viabilizar o start junto "
        "com o contrato de marketing já em andamento.",
        align="justify",
    )

    add_heading_styled(doc, "O que ficou alinhado vs. o que ficou pendente", 2)
    make_table(
        doc,
        ["Status", "Ponto"],
        [
            ["Alinhado", "Roteiro de 12 semanas e as seis fases (diagnóstico → handover)"],
            ["Alinhado", "Diagnóstico operacional presencial, com 18h provisionadas"],
            ["Alinhado", "Estruturação Comercial é custo único de implementação (não se repete)"],
            ["Alinhado", "Acompanhamento comercial após o projeto entra no executado de marketing, sem custo extra"],
            ["Alinhado", "Identidade do G6 Móvel segue o padrão visual da G6 Internet"],
            ["Pendente", "Validação da proposta e do investimento com Geraldo (Laura leva na terça)"],
            ["Pendente", "Kickoff com estrutura detalhada do plano (Denis)"],
            ["Pendente", "Reunião específica de lançamento móvel com Duarte / time de execução (Natã agenda)"],
        ],
        col_widths=[3.2, 13.3],
    )

    # ========== 2 ==========
    add_heading_styled(doc, "2. Contexto da G6 — o que foi diagnosticado", 1)
    add_p(
        doc,
        "Headline do deck: “Vocês não têm problema de lead. Têm um comercial reativo sem sistema.”",
        italic=True,
        color=RED,
    )
    add_p(
        doc,
        "A proposta foi montada a partir do contexto que Laura passou, das dores da operação e da "
        "movimentação do Mateus (novo gestor comercial, entrada em cerca de 30 dias). O mundo ideal "
        "apresentado: o gestor chega no meio do processo, treina com a V4 e herda uma máquina já "
        "estruturada — em vez de dividir tempo entre estruturar, trocar peça e bater meta.",
        align="justify",
    )

    add_heading_styled(doc, "Quatro falhas de sistema", 2)
    make_table(
        doc,
        ["Sintoma", "O que isso significa na G6"],
        [
            [
                "Entra ~1.000 / sai ~800",
                "Volume alto, retenção frágil. O crescimento líquido vaza todo mês. A batalha de aquisição não vira base.",
            ],
            [
                "CRM existe, mas não governa",
                "Sem motivo de perda, sem próximo passo, sem previsibilidade de pipeline. Dado solto. Não é painel de receita.",
            ],
            [
                "Meta vira comodidade",
                "Quando aperta, a meta desce para o time “bater” e receber comissão. Cultura de acomodação cascateada pela liderança.",
            ],
            [
                "Loja e PAP sem padrão",
                "5 internos + 2 porta a porta. Cada canal vende de um jeito, sem playbook, sem deparo, sem objeção mapeada.",
            ],
        ],
        col_widths=[5.0, 11.5],
    )

    add_p(doc, "Leitura do deck", size=12, bold=True, color=RED2, space_before=6)
    add_p(
        doc,
        "Demanda entra. Processo não segura. Receita escapa entre ativação, follow-up e retenção.",
        italic=True,
    )

    add_heading_styled(doc, "O que a operação já tem — e o que ainda não segura", 2)
    add_bullet(doc, "Time comercial de 7 pessoas, mais o Mateus na chegada.")
    add_bullet(doc, "Loja funcionando de um jeito e PAP de outro — nenhum dos dois com processo e meta claros.")
    add_bullet(doc, "Porta a porta sem deparo: região visitada, oportunidades mapeadas, o que trouxe, o que não trouxe, objeções.")
    add_bullet(doc, "Cliente oculto / mensagens de atendimento longas demais para um produto de alta intenção (primeiro plano de internet).")
    add_bullet(doc, "Base instalada sem jornada de upsell/upgrade — o time olha só “vendeu, acabou”.")
    add_bullet(doc, "Casos de “miséria”: cliente instalado que não pagou / não ativou — comissão apagada, receita nenhuma.")
    add_bullet(doc, "Cultura de atendimento forte (calor humano, proximidade). O comercial, porém, herdou o ritmo do atendimento e ficou reativo — falta a gana de meta sem perder qualidade de conversa.")
    add_bullet(doc, "Chegada do chip/móvel: se o time não estiver treinado para vender para a base, a G6 gasta de novo para ofertar um produto que a carteira já poderia absorver.")

    callout(
        doc,
        "Por que o diagnóstico precisa ser presencial",
        "Rafaela relatou que, de fora, a leitura era “ajuste de time”. Presencialmente, viu o perfil "
        "da casa: atendimento excelente, ritmo mais calmo, comercial seguindo a mesma cadência do "
        "atendimento. Sem estar in loco, o Denis estruturaria um plano genérico. A Fase 1 (18h) "
        "existe para ele ver cultura, resistência à mudança, loja, PAP, cidade e o dia a dia — "
        "não só entrevistar.",
    )

    # ========== 3 ==========
    add_heading_styled(doc, "3. A conta do vazamento — dados da operação e da reunião", 1)
    add_p(
        doc,
        "A G6 tem capacidade de vender. O que a operação não tem é capacidade de segurar — e de medir por que perde.",
        italic=True,
        color=RED,
    )

    add_heading_styled(doc, "3.1 Números do deck", 2)
    make_table(
        doc,
        ["Lado", "Indicador", "Dado"],
        [
            ["Entradas", "Ativações / mês", "~1.000 (número citado pela operação)"],
            ["Entradas", "Meta Forms (ano)", "1.269 leads  ·  CPL R$ 7,51"],
            ["Entradas", "WhatsApp", "Resposta em cerca de 1 minuto"],
            ["Saídas", "Cancelamentos jan–jun/2026", "3.755"],
            ["Saídas", "Janeiro", "415"],
            ["Saídas", "Fevereiro", "560"],
            ["Saídas", "Março", "614"],
            ["Saídas", "Abril", "746"],
            ["Saídas", "Maio", "715"],
            ["Saídas", "Junho", "705"],
        ],
        col_widths=[3.2, 5.5, 7.8],
    )

    callout(
        doc,
        "Pergunta que o deck deixou na mesa",
        "Qual é a taxa real de churn da G6? Para calcular, falta a carteira ativa no início de cada mês "
        "e o cancelamento por coorte, motivo, cidade e tempo de casa. Esse dado ainda não existe de "
        "forma conciliada — e sem ele, toda decisão de investimento é chute. A Estruturação Comercial "
        "instala exatamente essa inteligência no sistema.",
    )

    add_heading_styled(doc, "3.2 Conta de padaria apresentada na reunião (Rafaela → Laura / Geraldo)", 2)
    add_p(
        doc,
        "Esta conta não está no slide de vazamento do deck; foi falada na reunião para traduzir o "
        "impacto financeiro ao Geraldo. É estimativa de gestão, não conciliação contábil — e foi "
        "apresentada assim, de propósito, porque o tracking fino ainda não existe.",
        align="justify",
        italic=True,
        color=MUT,
    )
    make_table(
        doc,
        ["Premissa", "Número", "Leitura"],
        [
            ["CAC + instalação por cliente", "≈ R$ 800", "Custo médio para colocar o cliente dentro"],
            ["Payback mínimo", "8 meses", "Tempo para o cliente se pagar (zero a zero)"],
            ["Cancelamentos em 6 meses", "3.755", "Jan 415 → jun 705, série crescente no 1º semestre"],
            ["Investimento queimado (3.755 × R$ 800)", "≈ R$ 3 milhões", "Dinheiro já posto em aquisição + instalação"],
            ["Receita potencial deixada na mesa", "≈ R$ 4 milhões", "O que esses clientes deixariam ao ano"],
            ["Vazamento total estimado em 6 meses", "≈ R$ 7 milhões", "Investimento perdido + receita não realizada"],
            ["Contrato V4 marketing (12 meses)", "≈ R$ 134 mil", "12× R$ 11.180 — já fechado"],
            ["Meta de crescimento citada (Geraldo)", "5% ao ano", "Leitura: ~R$ 3 mi/ano a mais, ~R$ 300 mil/mês"],
        ],
        col_widths=[5.8, 3.8, 6.9],
    )

    add_heading_styled(doc, "3.3 O “mínimo que já paga o projeto” (cenário 2%)", 2)
    add_p(
        doc,
        "Rafaela ancorou o ROI sem superlativo: se a estruturação melhorar só 2% da perda — “que não é nada” "
        "frente ao que o método costuma mover — a G6 puxaria cerca de 75 clientes. Conta apresentada:",
        align="justify",
    )
    make_table(
        doc,
        ["Item", "Valor"],
        [
            ["Clientes retidos no cenário 2%", "≈ 75"],
            ["Custo já investido nesses 75 (não perdido)", "≈ R$ 60.000"],
            ["LTV / o que pagam em 12 meses", "≈ R$ 90.000"],
            ["Receita retida (60 + 90)", "≈ R$ 150.000"],
            ["Leitura", "Paga ~2 meses do contrato V4 + parte da Estruturação Comercial"],
        ],
        col_widths=[8.5, 8.0],
    )
    add_p(
        doc,
        "A tese: se a G6 não arruma o comercial agora, continua perdendo no longo prazo, porque o gasto "
        "para colocar o cliente dentro exige no mínimo 8 meses de permanência. Sem estrutura, nem se "
        "sabe quantos dos 415 que saíram em janeiro já tinham se pago.",
        align="justify",
    )

    # ========== 4 ==========
    add_heading_styled(doc, "4. Dores × o que a Estruturação Comercial fecha", 1)
    make_table(
        doc,
        ["Dor da G6", "O que o produto fecha"],
        [
            [
                "Processo comercial inconsistente — loja e PAP vendem sem padrão; não há jornada nem critério de etapa.",
                "Arquitetura dual-track loja × PAP, com critérios de etapa, papéis e jornada comercial alvo.",
            ],
            [
                "Cada vendedor vende de um jeito — performance depende de pessoa, não de método.",
                "Scripts, cadências, BPMN e matriz CHA. Treino com roleplay, não PDF na gaveta.",
            ],
            [
                "CRM não representa a realidade — pipeline não reflete ativação, perda nem follow-up.",
                "CRM como painel de receita: campos obrigatórios, motivos de perda, próxima ação, dashboard.",
            ],
            [
                "Gestor não conduz reunião comercial — Mateus precisa herdar rituais prontos.",
                "Rituais diário / semanal / pipeline review + onboarding do gestor na Fase 4 e 5.",
            ],
            [
                "Sem cadência e passagem de bastão — vendeu e acabou; ativação, save e retenção soltos.",
                "SLA marketing ↔ comercial ↔ operação; rotinas de ativação, follow-up e save no processo.",
            ],
            [
                "Treino não vira rotina — artefato sem instalação assistida regride.",
                "Fase 5 presencial (22h): imersão, roleplay, checklist semanal e correção de desvios.",
            ],
        ],
        col_widths=[8.25, 8.25],
    )

    # ========== 5 ==========
    add_heading_styled(doc, "5. Visão estratégica e promessa do projeto", 1)
    add_p(doc, "A ordem certa: fundação antes de pressão de meta.", italic=True, color=RED)

    make_table(
        doc,
        ["#", "Pilar", "O que instala"],
        [
            ["01", "Sistema — máquina comercial", "Processo, CRM e rituais transformam esforço individual em operação previsível."],
            ["02", "Gestão — gestor executa", "Com a casa estruturada, a liderança cobra comportamento e conversão — não improvisa processo."],
            ["03", "Receita — menos vazamento", "Ativação, follow-up e retenção entram no método. O CAC deixa de queimar no escuro."],
        ],
        col_widths=[1.6, 5.4, 9.5],
    )
    callout(
        doc,
        "Frase de ouro do deck",
        "Sem sistema, meta vira pressão. Com sistema, meta vira consequência.",
    )

    add_heading_styled(doc, "Promessa", 2)
    add_p(
        doc,
        "Em um projeto fechado, a V4 constrói e instala a fundação comercial da G6. "
        "Processo, CRM em uso, scripts, cadências, indicadores e rituais de gestão — "
        "operando de verdade pelo time de loja e PAP.",
        align="justify",
    )
    add_p(doc, "Não é entregar PDF. É deixar a operação rodando.", bold=True, color=RED)

    add_p(
        doc,
        "Na reunião, Rafaela reforçou o mesmo ponto: a ideia não é auditar, entregar um documento "
        "e deixar a G6 aplicar. É auditar, implementar, treinar, criar rotina, garantir que está "
        "rodando para o Mateus — e só então passar o bastão.",
        align="justify",
    )

    # ========== 6 ==========
    add_heading_styled(doc, "6. Formato, escopo e o que fica instalado", 1)
    make_table(
        doc,
        ["Bloco", "Definição"],
        [
            ["Duração", "12 semanas — diagnosticar, construir, instalar e fazer handover com adesão."],
            ["Carga V4", "106,5h (média/alta): análise, desenho, CRM, QA e instalação — não só calls. +6h de gestão de projeto, QA e alinhamentos internos."],
            ["Execução", "V4 constrói e conduz a instalação; o time G6 opera no dia a dia."],
            ["Ritmo", "Encontros semanais com sponsor/gestor e checkpoints com o time comercial (Denis organiza)."],
            ["Presencial", "Fase 1 (diagnóstico) e Fase 5 (instalação/treino) — onde a adesão se decide."],
            ["Saída", "Rotina de gestão instalada + plano de continuidade de 60 dias."],
            ["Janela operacional", "Setembro–novembro de execução + plano de 60 dias; leitura de cobertura até janeiro, com Mateus já operando a máquina."],
        ],
        col_widths=[3.6, 12.9],
    )

    add_heading_styled(doc, "Encaixe com a chegada do Mateus", 2)
    add_p(
        doc,
        "Primeiros ~30 dias: Fases 1–2 (e início do CRM) avançam antes do gestor entrar. "
        "Quando o Mateus chega, o projeto já está em treinamento do time e dele. "
        "O handover (passagem de bastão do Denis) coincide com o fim do onboarding interno. "
        "O treinamento do gestor não vira problema da diretoria — é troca Denis ↔ Mateus.",
        align="justify",
    )

    add_heading_styled(doc, "O que o produto instala", 2)
    make_table(
        doc,
        ["Fundação operacional", "Controle e adesão"],
        [
            ["Arquitetura loja × PAP com critérios de etapa", "CRM como painel de receita (não agenda)"],
            ["Scripts, cadência e papéis claros por canal", "Higiene, motivos de perda e dashboard de uso"],
            ["Rituais de gestão (diária / semanal / pipeline review)", "Treinamento com instalação assistida + roleplay"],
            ["Onboarding operacional de vendedores", "Rotinas de ativação, follow-up e save no processo"],
        ],
        col_widths=[8.25, 8.25],
    )
    add_p(doc, "O produto não vende documento. Vende capacidade instalada no time.", italic=True, color=RED)

    # ========== 7 ==========
    add_heading_styled(doc, "7. As seis fases (com entregas e encaixe G6)", 1)
    add_p(
        doc,
        "Seis fases. Uma fundação comercial instalada. Total: 100,5h de fases + 6h de gestão/QA = 106,5h.",
        align="justify",
    )
    make_table(
        doc,
        ["Fase", "Horas", "Nome", "O que resolve"],
        [
            ["01", "18h", "Diagnóstico operacional", "Operação real, não documento bonito para o problema errado. Presencial."],
            ["02", "15h", "Arquitetura comercial", "Sistema operacional que o time consegue executar."],
            ["03", "15,5h", "CRM e controle", "Em paralelo à F2 — método vira workflow. Brian (CRM) entra aqui."],
            ["04", "20h", "Camada estratégica", "Playbook, comissão e breakeven. ~9ª semana / chegada do Mateus."],
            ["05", "22h", "Instalação da rotina", "Adesão: imersão, roleplay e rituais assistidos. Presencial."],
            ["06", "10h", "Handover 60 dias", "Continuidade + retenção + riscos de regressão."],
        ],
        col_widths=[1.6, 2.0, 4.6, 8.3],
    )

    # Fase 1
    add_heading_styled(doc, "Fase 01 — Diagnóstico operacional  ·  18h  ·  presencial", 2)
    add_p(
        doc,
        "Entender a operação real da G6 antes de desenhar qualquer playbook. Não é só entrevista: "
        "é assistir a operação. 18 horas provisionadas para conversar com Laura, com a gestão, "
        "com o time de loja e com o PAP — o suficiente para mapear a rotina e os pontos da equação "
        "(onde 1+1 está dando zero).",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "Diagnóstico comercial detalhado",
        "Entrevistas com sponsor, gestor, top e bottom performer",
        "Auditoria do CRM / RD atual",
        "Mapa de gargalos por etapa",
        "Priorização das dores que impedem venda previsível",
    ]:
        add_bullet(doc, t)
    add_p(doc, "Encaixe G6: presencial no onboarding — time fechado, loja + PAP.", italic=True, color=MUT)

    # Fase 2
    add_heading_styled(doc, "Fase 02 — Arquitetura operacional comercial  ·  15h", 2)
    add_p(
        doc,
        "Denis volta da imersão e constrói o plano de melhoria a partir da auditoria: o que a G6 "
        "acerta e o que precisa correção, para um comercial mais fluido, com menos “expresso” e "
        "mais passagem de jornada.",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "ICP / segmentos e jornada comercial alvo",
        "Pipeline e critérios de etapa · papéis",
        "SLA marketing ↔ comercial ↔ gestor (e comercial → financeiro/operação)",
        "Metas de atividade e indicadores",
        "Scripts · cadências · BPMN · matriz CHA",
    ]:
        add_bullet(doc, t)
    add_p(doc, "Encaixe G6: dual-track loja × PAP.", italic=True, color=MUT)

    # Fase 3
    add_heading_styled(doc, "Fase 03 — CRM e controle  ·  15,5h  ·  paralelo à Fase 2", 2)
    add_p(
        doc,
        "O analista de CRM (Brian) entra com a análise do Denis para reorganizar — ou, se a Fase 1 "
        "concluir que o CRM atual não é o ideal, implementar outro (licenças por conta da G6; "
        "configuração e implementação inclusas). Faturamento é a última etapa: cada etapa da jornada "
        "precisa ter métrica de sucesso (analogia do bolo / 250 g de farinha).",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "Reorganização do pipeline no CRM",
        "Campos obrigatórios · motivos de perda",
        "Atividades obrigatórias e próximas ações",
        "Dashboard de gestão e conversão (o “monitor” do gestor — Mateus)",
        "Regras de higiene · rotina de registro",
    ]:
        add_bullet(doc, t)
    add_p(
        doc,
        "Encaixe G6: origem, tempo de resposta, status, dor, capacidade, produto, próxima ação e motivo de perda.",
        italic=True,
        color=MUT,
    )

    # Fase 4
    add_heading_styled(doc, "Fase 04 — Camada estratégica  ·  20h", 2)
    add_p(
        doc,
        "Processo, gestão, metas, incentivo e viabilidade financeira. Playbook apresentado à diretoria "
        "e à gestão; política de comissionamento; breakeven da operação (“se eu inputar X de marketing "
        "com este time e este formato, aonde eu chego?”). Perspectiva: por volta da 9ª semana, com o "
        "Mateus já em onboarding — ele é onboardado na filosofia do projeto, não no jeito antigo.",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "Playbook comercial orientado ao uso (não à gaveta)",
        "Política de comissionamento",
        "Breakeven da operação comercial",
    ]:
        add_bullet(doc, t)
    add_p(
        doc,
        "Encaixe G6: comissão loja ≠ PAP. Meta como piso. Breakeven amarra CAC, payback e retenção.",
        italic=True,
        color=MUT,
    )

    # Fase 5
    add_heading_styled(doc, "Fase 05 — Instalação da rotina  ·  22h  ·  presencial", 2)
    add_p(
        doc,
        "Volta presencial a Minas. É a fase que garante que os artefatos sejam usados. Denis conduz "
        "o comercial; Brian conduz o treino de CRM. Roleplay na loja e no PAP. Primeiras rotinas de "
        "pipeline review. Checklist semanal para correção de desvios. Se amanhã a G6 contratar mais "
        "10 vendedores, já existe o pacote de recebimento, treino e documentos.",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "Onboarding operacional de vendedores",
        "Imersão do time + treinamento do gestor no CRM",
        "Roleplay loja e PAP",
        "Pipeline review e reunião semanal assistidas",
        "Checklist semanal · correção de desvios",
    ]:
        add_bullet(doc, t)
    add_p(doc, "Encaixe G6: imersão presencial + roleplay. Mateus herda a rotina.", italic=True, color=MUT)

    # Fase 6
    add_heading_styled(doc, "Fase 06 — Handover e continuidade  ·  10h", 2)
    add_p(
        doc,
        "Últimos 7 a 10 dias. Transferência da propriedade intelectual: tudo documentado, editado e "
        "corrigido conforme a G6 pediu. Score de maturidade antes/depois. Rotinas de retenção e "
        "expansão (o escoamento da base). Mapa de risco de regressão (“se voltarmos a fazer o que "
        "fazíamos, o que quebra?”). Plano tático dos 60 dias seguintes — cobertura de operação até "
        "janeiro; a partir daí, replicar o processo.",
        align="justify",
    )
    add_p(doc, "Entregas", bold=True, color=RED2, space_after=4)
    for t in [
        "Reunião final de transferência",
        "Score de maturidade antes/depois",
        "Plano de 60 dias",
        "Rotinas de retenção e expansão",
        "Riscos de regressão · evolução",
    ]:
        add_bullet(doc, t)
    add_p(doc, "Encaixe G6: motivos de perda, save/upsell e ritual de gestão.", italic=True, color=MUT)

    # ========== 8 ==========
    add_heading_styled(doc, "8. Três pilares: sistemas, processos e pessoas", 1)
    add_p(
        doc,
        "Denis Orosco — consultor de receita da V4, responsável pela frente de produtos Sales da unidade — "
        "estruturou a fala em três pilares. Essa camada não está como slide separado no deck, mas foi "
        "o fio da apresentação das fases.",
        align="justify",
    )
    make_table(
        doc,
        ["Pilar", "Pergunta que o projeto responde", "Na prática G6"],
        [
            [
                "Sistemas",
                "O CRM atual serve? Vale manter ou trocar? Qual a melhor arquitetura da stack?",
                "Governança do RD/CRM, campos, motivos de perda, dashboard tático do Mateus. Se precisar trocar, implementação inclusa.",
            ],
            [
                "Processos",
                "Como os fluxos se conectam? Onde está a passagem de bastão venda → operação?",
                "SLA, jornada loja × PAP, cadência, script, higiene, rituais semanais, checklist de desvio.",
            ],
            [
                "Pessoas",
                "Como capacito o time, a nova liderança e a direção para gerir o novo método?",
                "Imersão, roleplay, treino do Mateus, treino da direção na cobrança tático-operacional.",
            ],
        ],
        col_widths=[3.0, 6.5, 7.0],
    )

    add_heading_styled(doc, "Credenciais apresentadas (Denis)", 2)
    add_bullet(doc, "12 anos em áreas que conectam marketing e comercial.")
    add_bullet(doc, "Google Expert 2025 (um dos 170 profissionais certificados no país) — linha forte de gestão de tráfego, para o comercial não ser atrito do investimento de mídia.")
    add_bullet(doc, "Lean / Six Sigma: Black Belt certificado (único player da V4 Company com essa formação no momento da reunião); rota para Master Black Belt. Mapeamento com roteiro DMAIC.")
    add_bullet(doc, "Win by Design (Vale do Silício) — certificação em arquitetura de receita; menos de 300 profissionais certificados no Brasil.")
    add_bullet(doc, "2 anos na V4 (24º mês em agosto/2026). 1º ano em estruturas complexas — conta de R$ 1,4 bi/ano, 3.200 vendedores, 260 gerentes, 6 estados. Desde setembro anterior na frente dos produtos Sales (produto prototipado por ele).")
    add_p(
        doc,
        "Postura na conta: mediador. A máquina a G6 executa; a V4 ajuda a achar o melhor caminho e instala.",
        italic=True,
        color=MUT,
    )

    # ========== 9 ==========
    add_heading_styled(doc, "9. Critérios de Receita Previsível", 1)
    add_p(
        doc,
        "Receita Previsível não é slogan. São os critérios observáveis para dizer que a fundação está instalada:",
        align="justify",
    )
    make_table(
        doc,
        ["#", "Critério de sucesso"],
        [
            ["1", "Pipeline configurado e usado — não só “existente”"],
            ["2", "Reunião semanal acontecendo de verdade"],
            ["3", "Follow-up sendo executado"],
            ["4", "Gestor usando indicadores"],
            ["5", "Time nos scripts e critérios"],
            ["6", "Onboarding documentado"],
            ["7", "Rotinas de retenção em operação"],
            ["8", "Sponsor com previsibilidade"],
        ],
        col_widths=[1.6, 14.9],
    )

    add_heading_styled(doc, "Recap do deck — o que está em jogo", 2)
    make_table(
        doc,
        ["Dores", "Dados"],
        [
            ["Comercial reativo: entra ~1.000 / base quase não cresce", "~1.000 ativações/mês · Meta Forms 1.269 leads · CPL R$ 7,51"],
            ["3.755 cancelamentos em 6 meses — sem churn real conciliado", "Saídas: 415 · 560 · 614 · 746 · 715 · 705 (jan–jun)"],
            ["CRM sem governança — não é painel de receita", "WhatsApp responde em ~1 minuto"],
            ["Loja e PAP sem padrão, script nem cadência", "5 vendedores loja + 2 PAP · Mateus em ~30 dias"],
            ["Ativação, save e retenção fora do processo", "Projeto: 12 semanas · 106,5h · dual-track loja × PAP"],
        ],
        col_widths=[8.25, 8.25],
    )
    callout(
        doc,
        "Fechamento do recap",
        "A Estruturação Comercial instala a máquina que fecha essas dores — com método, adesão e gestão.",
    )

    # ========== 10 ==========
    add_heading_styled(doc, "10. Investimento, condição comercial e viabilidade", 1)

    price = doc.add_table(rows=1, cols=2)
    set_table_width(price, 16.5)
    left, right = price.cell(0, 0), price.cell(0, 1)
    shade_cell(left, DARK)
    shade_cell(right, SOFT)
    set_cell_borders(left, DARK, "4")
    set_cell_borders(right, RED, "8")
    set_cell_margins(left, 160, 160, 160, 160)
    set_cell_margins(right, 140, 140, 140, 140)
    left.text = ""
    p = left.paragraphs[0]
    r = p.add_run("VALOR DE REFERÊNCIA DO PRODUTO")
    set_run_font(r, size=8, bold=True, color="F5B5B8")
    p2 = left.add_paragraph()
    r = p2.add_run("R$ 89.635,85")
    set_run_font(r, size=16, color="C9A0A0")
    p2.runs[0].font.strike = True
    p3 = left.add_paragraph()
    r = p3.add_run("CONDIÇÃO COMERCIAL G6  ·  BASE ATIVA")
    set_run_font(r, size=8, bold=True, color="F5B5B8")
    p4 = left.add_paragraph()
    r = p4.add_run("R$ 71.582,32")
    set_run_font(r, size=26, bold=True, color=WHITE)
    p5 = left.add_paragraph()
    r = p5.add_run("12 semanas · 106,5h · implementação assistida\ndual-track loja × PAP · presenciais nas fases críticas")
    set_run_font(r, size=10, color="E8D0D0")

    right.text = ""
    p = right.paragraphs[0]
    r = p.add_run("FORMA APRESENTADA NA REUNIÃO")
    set_run_font(r, size=8, bold=True, color=RED)
    add_cell_run(right, "12 recorrências  ·  R$ 5.965,19 / mês", size=14, bold=True, color=INK)
    add_cell_run(right, "Pagamento recorrente (sem cartão de crédito), para não competir com a liberação do cartão do contrato de marketing.", size=10, color=MUT)
    add_cell_run(right, "Redução de R$ 18.053,53 em relação à referência, por Laura ser cliente da base ativa.", size=10, color=MUT)
    add_cell_run(right, "Custo único de implementação. Não se repete no ano 2. O acompanhamento do comercial segue no executado de marketing, sem custo adicional.", size=10, color=MUT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    add_heading_styled(doc, "O que está incluso / por que agora / resultado esperado", 2)
    make_table(
        doc,
        ["Bloco", "Conteúdo"],
        [
            ["Incluso", "Capacidade instalada: diagnóstico → arquitetura → CRM como painel de receita → estratégica → instalação → handover 60 dias."],
            ["Por que agora", "Parar o vazamento. Com 3.755 saídas em 6 meses e churn real ainda não conciliado, estruturar o comercial é a decisão que dá previsibilidade."],
            ["Resultado esperado", "Máquina comercial: processo, rituais e gestão rodando — para vender, ativar e reter com método."],
        ],
        col_widths=[3.6, 12.9],
    )

    add_heading_styled(doc, "Como a conta fecha no ano (para o Geraldo)", 2)
    add_p(
        doc,
        "Na reunião, Laura somou o marketing já fechado com a Estruturação Comercial. Rafaela pediu "
        "para levar ao Geraldo o depara com os R$ 7 milhões — senão ele olha preço, não investimento.",
        align="justify",
    )
    make_table(
        doc,
        ["Linha", "Base de cálculo", "Total"],
        [
            ["Marketing V4 (já fechado)", "12 × R$ 11.180", "R$ 134.160"],
            ["Estruturação Comercial (proposta)", "R$ 71.582,32  (ou 12 × R$ 5.965,19)", "R$ 71.582,32"],
            ["Pacote anual marketing + EC", "Laura arredondou o mensal para ~R$ 17.180", "≈ R$ 205.700  (Rafaela citou R$ 204 mil)"],
            ["+ mídia (~R$ 10 mil/mês)", "R$ 120.000 no ano", "≈ R$ 324.000 no ano"],
            ["Vazamento estimado em 6 meses", "Conta de padaria da reunião", "≈ R$ 7.000.000"],
            ["Relação investimento × vazamento", "324 mil / 7 milhões", "Menos de 5%  (Rafaela: “nem 10%”)"],
        ],
        col_widths=[5.2, 6.3, 5.0],
    )

    callout(
        doc,
        "Argumento de decisão levado ao Geraldo",
        "Ou a G6 mexe na máquina agora para evitar outro semestre de vazamento, ou continua sentada "
        "esperando o comercial mudar sozinho. Daqui a 3 meses a V4 entrega lead qualificado e o furo "
        "do balde continua do lado comercial — e a pauta vira briga de conversão, não expansão. "
        "O investimento para buscar receita previsível está mais barato do que um mês perdendo ~800 clientes.",
    )

    add_heading_styled(doc, "O que é único vs. o que renova", 2)
    add_bullet(doc, "Estruturação Comercial: custo único. Implementa, o executado acompanha, não gera nova fatura dessa frente.", bold_prefix="")
    add_bullet(doc, "Marketing (execução): 12 meses. Depois a G6 renova ou não. Mídia, criativo e operação de growth estão nesse bloco.")
    add_bullet(doc, "Estruturação estratégica prévia (plano de marketing já feito): também é etapa única de entrada — auditar e planejar antes de executar, para não repetir erro.")
    add_p(
        doc,
        "Rafaela não negocia valor na mesa de estratégia. Dinheiro fica com a Sara. Na reunião, o "
        "pedido explícito a Laura: transparência se o número não couber — “calculamos que se fosse X "
        "a gente conseguiria” — para a V4 tentar viabilizar o start agora, não daqui a alguns meses.",
        align="justify",
    )

    # ========== 11 ==========
    add_heading_styled(doc, "11. Ata da reunião de 14/08/2026", 1)
    add_p(
        doc,
        "Fonte: observações Gemini + transcrição da call “[Comercial] V4 Company :::: G6 Internet”, "
        "14 de agosto de 2026. Duração aproximada: 1h23. Gravação anexa no Google Docs.",
        italic=True,
        color=MUT,
    )

    add_heading_styled(doc, "Participantes e papéis", 2)
    make_table(
        doc,
        ["Pessoa", "Papel na reunião"],
        [
            ["Laurelise Santos (Laura)", "Sponsor G6. Precisa validar com Geraldo. Time comercial de 7 pessoas + Mateus."],
            ["Rafaela Tolomeu Messias", "Conduziu a leitura estratégica, a conta de vazamento e o investimento."],
            ["Sara Pizzico", "Comercial V4. Dona da negociação de valor (não esteve na fala de preço; Rafaela redireciona para ela)."],
            ["Denis Orosco", "Consultor de receita / Sales. Detalhou fases, credenciais e formato. Saiu ~1h07 por outra agenda."],
            ["Geraldo (ausente)", "Sócio. Dono da validação financeira. Laura conversa com ele na terça."],
            ["Mateus (ausente)", "Novo gestor comercial. Entra em ~30 dias. Herda a máquina."],
            ["Brian (citado)", "Analista de CRM. Entra na Fase 3 e no treino da Fase 5. Duarte já o apresentou à G6."],
            ["Duarte / Natã (citados)", "Execução de marketing / agenda. Pauta de lançamento móvel na semana seguinte."],
        ],
        col_widths=[5.0, 11.5],
    )

    add_heading_styled(doc, "Resumo da ata (Gemini)", 2)
    add_p(
        doc,
        "A reunião abordou a reestruturação comercial para estancar perdas financeiras e planejar o "
        "lançamento de novos serviços. A organização identificou falhas graves na governança de sistemas "
        "e processos, com prejuízo acumulado estimado de R$ 7 milhões em 6 meses. Decidiu-se o plano de "
        "12 semanas focado em sistemas, processos e capacitação. A viabilidade financeira do projeto foi "
        "contextualizada; o planejamento do serviço móvel foi iniciado como pauta paralela.",
        align="justify",
    )

    add_heading_styled(doc, "Decisões registradas", 2)
    make_table(
        doc,
        ["Tipo", "Decisão"],
        [
            ["Precisa de mais conversa", "Validação da proposta de reestruturação comercial com o sócio Geraldo antes do fechamento definitivo."],
            ["Alinhada", "Adoção do roteiro de 12 semanas da V4, incluindo as seis fases de diagnóstico, arquitetura, implementação e treinamento."],
            ["Alinhada", "Diagnóstico operacional in loco, com 18 horas provisionadas para mapeamento."],
            ["Alinhada", "Identidade visual do G6 Móvel segue o mesmo padrão da G6 Internet, para credibilidade e reconhecimento."],
        ],
        col_widths=[4.4, 12.1],
    )

    add_heading_styled(doc, "Próximas etapas (com dono)", 2)
    make_table(
        doc,
        ["Dono", "Ação"],
        [
            ["Denis Orosco", "Realizar o diagnóstico operacional in loco (18h): mapear rotinas, entrevistar gestão, observar loja e campo."],
            ["Denis Orosco", "Apresentar a estrutura detalhada do plano no kickoff."],
            ["Denis Orosco", "Desenvolver a arquitetura operacional e comercial: metas, indicadores, scripts e cadências."],
            ["Denis Orosco", "Treinar o time comercial e o gestor Mateus na nova estrutura e na usabilidade do CRM."],
            ["Laura", "Apresentar a proposta ao Geraldo e demais sócios; validar investimento e aceitação do projeto (terça)."],
            ["Rafaela", "Briefing ao Duarte sobre o lançamento G6 Móvel: PDV, abrangência inicial, nome e comunicação."],
            ["Rafaela / Natã", "Agendar reunião na semana seguinte com o time de execução para lançamento do produto."],
        ],
        col_widths=[4.0, 12.5],
    )

    add_heading_styled(doc, "Linha do tempo da conversa (pauta comercial)", 2)
    make_table(
        doc,
        ["Min", "O que aconteceu"],
        [
            ["04:34", "Rafaela abre o contexto: dores, Mateus, receita minando, dados soltos. Objetivo: entregar máquina pronta ao gestor."],
            ["06:46", "Espelho: sem receita previsível; entra ~1.000 / sai ~800; CRM sem governança; meta que desce; loja e PAP sem processo."],
            ["09:52", "Conta do vazamento: R$ 800, 8 meses, 3.755 cancelamentos, ~R$ 3 mi + ~R$ 4 mi = ~R$ 7 mi em 6 meses."],
            ["12:20", "Cenário 2% (~75 clientes, ~R$ 150 mil retidos). Sem estrutura, o balde continua furado."],
            ["16:54", "Denis se apresenta como consultor de receita / mediador da máquina Sales."],
            ["19:27", "Jornada: mensagens longas, sem upsell na base, risco na chegada do chip."],
            ["21:26", "Promessa do projeto fechado (não é PDF). 12 semanas, 106h, F1 e F5 presenciais, handover 60 dias."],
            ["26:15", "Credenciais Denis: Google Expert, Black Belt, Win by Design, conta de R$ 1,4 bi."],
            ["31:41", "Três pilares: sistemas, processos, pessoas. Roadmap das 6 fases."],
            ["34:39", "Por que presencial: cultura reativa vista in loco. 18h de diagnóstico."],
            ["40:22", "Fases 2 e 3: arquitetura + CRM (Brian), campos, dashboard, métricas por etapa."],
            ["54:37", "Fase 4: playbook, comissão, breakeven, onboarding do Mateus ~9ª semana."],
            ["55:44", "Fase 5: volta a Minas, imersão, roleplay, rituais, checklist."],
            ["57:43", "Fase 6: transferência de IP, maturidade, retenção, risco de regressão, plano 60 dias."],
            ["59:49", "Laura confirma encaixe com a dor da quarta-feira. Quer fazer; precisa dos sócios. Antecipa resistência do time de 7 + Mateus."],
            ["01:02", "Rafaela: marketing + EC < 10% dos R$ 7 mi. Valor não se negocia com ela — vai para a Sara."],
            ["01:05", "Investimento: de R$ 89.635,85 para R$ 71.582,32; 12× recorrente, sem cartão."],
            ["01:07", "Denis se despede. Laura: precisa validar; soma ~R$ 17 mil/mês com o marketing. Rafaela ancora R$ 204 mil vs. R$ 7 mi."],
            ["01:11", "EC é custo único; acompanhamento já está no executado. Marketing renova ou não aos 12 meses."],
            ["01:13", "Laura fala com Geraldo na terça e puxa a V4 se precisar. Rafaela pede transparência para viabilizar."],
            ["01:15", "Pauta extra: lançamento móvel (nome, PDV, Piranga Sul / Tejubá, mesma cara da G6)."],
        ],
        col_widths=[2.2, 14.3],
    )

    add_heading_styled(doc, "Posição da Laura na mesa", 2)
    add_p(
        doc,
        "Laura confirmou que o escopo responde à pergunta que ela fez na reunião anterior (“vocês vão "
        "olhar o comercial também?”). Enxerga a mudança como corte de era: a G6 não é mais a empresa "
        "pequena em que se conhecia cada funcionário. Já prevê que parte do time de 7 “espanne”, e que "
        "a chegada do Mateus sozinha já gera atrito. Assume a mudança porque “não dá para ser pai e mãe "
        "o tempo todo”. Sentimento declarado: “quero fazer” — com o caveat de repassar aos sócios. "
        "O valor não tinha sido pensado antes; surgiu da conversa daquela semana.",
        align="justify",
    )

    # ========== 12 ==========
    add_heading_styled(doc, "12. Pauta complementar: lançamento G6 Móvel / Chip", 1)
    add_p(
        doc,
        "Depois da proposta, Laura pediu ajuda no lançamento do serviço móvel, que deve sair mais "
        "rápido do que o previsto. Esta pauta não faz parte do escopo da Estruturação Comercial — "
        "foi encaminhada para o time de execução (Duarte) em reunião específica na semana seguinte. "
        "Entra neste documento porque foi apresentada na mesma call e se conecta com a tese comercial: "
        "sem time treinado para vender na base, o chip vira novo CAC em cima de carteira já paga.",
        align="justify",
    )
    make_table(
        doc,
        ["Tema", "O que ficou"],
        [
            ["Nome em avaliação", "G6 Chip  ·  G6 Móvel  ·  G6 Celular  ·  G6 5G"],
            ["Linha de comunicação sugerida (Laura)", "“A qualidade e a constância que você conhece na sua casa, agora no seu bolso.”"],
            ["Marca", "Mesma cara da G6 Internet. Credibilidade de mercado (analogia: a Vivo não muda de cara entre casa e celular)."],
            ["Praça 1", "Piranga Sul (transcrição: “Pirango Sul”) — ~6.000 habitantes. Lançamento pequeno para validar."],
            ["Praça 2", "Tejubá (transcrição) — ~90 mil habitantes, dos quais ~20 mil já são clientes G6. Venda nova + base."],
            ["Materiais de PDV", "Wind banner na porta da loja; totem/display de papelão; wobblers (“bolachinha”). Produzir o kit para todas as lojas mesmo usando só a primeira agora."],
            ["CRM / base", "Neste momento Laura não precisa de disparo massivo na base toda. Validar off na cidade pequena; campanha/disparo só daquela região se fizer sentido."],
            ["Próximo passo", "Rafaela passa o briefing ao Duarte e pede ao Natã a agenda da semana seguinte."],
        ],
        col_widths=[4.4, 12.1],
    )

    # ========== 13 ==========
    add_heading_styled(doc, "13. Como levar este caso ao Geraldo", 1)
    add_p(
        doc,
        "Roteiro que a Rafaela pediu para Laura usar — útil também para a Sara no follow-up:",
        align="justify",
    )
    add_bullet(doc, "Não abrir pelo preço. Abrir pelos R$ 7 milhões em 6 meses e pelos 3.755 cancelamentos.")
    add_bullet(doc, "Mostrar que cada cliente precisa de 8 meses para se pagar (CAC + instalação ≈ R$ 800).")
    add_bullet(doc, "Deparar os ~R$ 204 mil no ano (marketing + EC) — ou ~R$ 324 mil com mídia — com o vazamento. É menos de 10%.")
    add_bullet(doc, "Deixar claro: Estruturação Comercial é custo único; não vira mensalidade eterna.")
    add_bullet(doc, "Explicar o timing do Mateus: os primeiros 30 dias estruturam a casa; ele herda máquina + treino, em vez de improvisar processo.")
    add_bullet(doc, "Cenário mínimo (2%): ~R$ 150 mil retidos já ajudam a pagar o projeto. Qualquer melhoria real tende a ser maior.")
    add_bullet(doc, "Risco de esperar 3 meses: a V4 entrega lead, o comercial continua furado, a relação vira cobrança de conversão.")
    add_bullet(doc, "Se o número não couber, voltar com o teto que cabe — a V4 tenta viabilizar o start agora (recorrência já foi o gesto para não travar no cartão).")

    callout(
        doc,
        "Frase de fechamento (script interno)",
        "O que a gente propõe não é mais esforço comercial. É previsibilidade: parar o vazamento com "
        "método instalado na loja e no PAP — em doze semanas.",
    )

    # ========== 14 ==========
    add_heading_styled(doc, "14. Fontes e materiais de apoio", 1)
    make_table(
        doc,
        ["Material", "Onde"],
        [
            ["Ata e transcrição da reunião 14/08/2026", "https://docs.google.com/document/d/12M9kFqxpMApDL0yZMOgtmz1UbBXQvFDMvXZCsnZZSAE/edit"],
            ["Deck Estruturação Comercial (18 slides)", "proposta-g6-estruturacao-comercial.html  ·  /g6/  ·  /g6-ec/"],
            ["PDF do deck", "G6-Internet_Estruturacao-Comercial_V4.pdf"],
            ["Script interno de narrativa (slide a slide)", "script-narrativa-g6-estruturacao-comercial.docx"],
            ["Proposta Growth + CRM (12 meses, contexto anterior)", "proposta-g6-internet.html"],
        ],
        col_widths=[6.5, 10.0],
    )

    add_p(
        doc,
        "Documento consolidado pela V4 Company a partir do deck apresentado e da reunião comercial "
        "de 14 de agosto de 2026. Valores e prazos conforme falados na mesa e registrados no deck. "
        "A conta de R$ 7 milhões é estimativa de gestão (não conciliação contábil), e foi apresentada "
        "assim porque o churn real ainda não está conciliado — exatamente o que o projeto se propõe a instalar.",
        size=9,
        color=MUT,
        italic=True,
        space_before=12,
    )
    add_p(
        doc,
        "G6 Internet × V4 Company  ·  Estruturação Comercial  ·  Confidencial",
        size=9,
        color=RED,
        bold=True,
        align="center",
        space_before=8,
    )

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
